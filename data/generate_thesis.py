"""
SmartClaim AI Thesis Generator
Generates a full thesis .docx mirroring the Jeddah Beacon structure.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_body(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def create_thesis():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)

    # ============================================================
    # TITLE PAGE 1
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Title")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SmartClaim AI: A High-Granularity Hybrid Expert Pipeline for Saudi Vehicle Accident Cost Prediction from Arabic Text")
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ra'ana Hatim Shaikh")
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("by:")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A thesis submitted for the requirements of the degree of\nProfessional Master in Artificial Intelligence")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Faculty of Computing and Information Technology\nKing Abdulaziz University")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Jeddah, Saudi Arabia")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dhu'l-Hijjah 1447 H - June 2026 G")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ============================================================
    # TITLE PAGE 2 (with Advisor)
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Title")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SmartClaim AI: A High-Granularity Hybrid Expert Pipeline for Saudi Vehicle Accident Cost Prediction from Arabic Text")
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ra'ana Hatim Shaikh")
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("by:")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A thesis submitted for the requirements of the degree of\nProfessional Master in Artificial Intelligence")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Advisor:")
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dr. Somayah Albaradei")
    run.font.size = Pt(14)
    
    doc.add_page_break()
    
    # ============================================================
    # TITLE PAGE 3 (Committee Page)
    # ============================================================
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Title")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SmartClaim AI: A High-Granularity Hybrid Expert Pipeline for Saudi Vehicle Accident Cost Prediction from Arabic Text")
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ra'ana Hatim Shaikh")
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("by:")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A thesis submitted for the requirements of the degree of\nProfessional Master in Artificial Intelligence")
    
    doc.add_paragraph()
    
    # Committee table
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Name", "Rank", "Field", "Signature"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(cell, "D9E2F3")
    
    roles = ["Advisor", "Co-Advisor", "Internal Examiner", "External Examiner"]
    for i, role in enumerate(roles):
        table.rows[i+1].cells[0].text = role
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Faculty of Computing and Information Technology\nKing Abdulaziz University")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Jeddah, Saudi Arabia")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dhu'l-Hijjah 1447 H - June 2026 G")
    
    doc.add_page_break()
    
    # ============================================================
    # COPYRIGHT
    # ============================================================
    add_heading_styled(doc, "Copyright", level=1)
    add_body(doc, "All rights reserved to King Abdulaziz University. It is not permitted to copy or reissue this scientific thesis or any part of it in any way or by any means except with the prior written permission from the author or the scientific department. It is also not allowed to translate it into any other language and it is necessary to refer to it when citing. This page must be part of any additional copies.")
    
    doc.add_page_break()
    
    # ============================================================
    # DEDICATION
    # ============================================================
    add_heading_styled(doc, "Dedication", level=1)
    
    add_body(doc, "To those whose quiet belief carried me farther than I ever imagined I could go\u2026", italic=True)
    
    add_body(doc, "To my family, who stood behind every step I took and lived within every achievement I reached. This journey was never a solitary one; it was ours together, carried with everything it held of exhaustion, hope, waiting, and joy.")
    
    add_body(doc, "On the nights the hours grew long and the drive grew dim, it was you who kept me going \u2014 quiet prayers offered in the dark, words spoken at exactly the right moment, and a presence that never needed an invitation.")
    
    add_body(doc, "To those who taught me from the very beginning that knowledge is an honour, that dedication never fails its owner, and that a person is shaped by those around them before anything within them. You were my first school, my deepest lesson, and the truest meaning behind everything I have achieved.")
    
    add_body(doc, "You were my roots when I needed steadiness, my wings when I needed to rise, and my refuge when I needed peace. This work would never have seen the light of day without everything you gave in silence, and everything you carried with love.")
    
    add_body(doc, "And to myself \u2014 to that version of me that refused to give up when the horizon felt impossibly narrow and refused to turn back when the weight grew heavy. To every late night spent in uncertainty, not knowing whether this would ever be finished \u2014 and to every moment of self-doubt that was met, somehow, with the decision to keep going anyway.")
    
    add_body(doc, "This achievement is mine too \u2014 and with everything this work carries of effort, hope, and gratitude, I dedicate this work.")
    
    doc.add_page_break()
    
    # ============================================================
    # ACKNOWLEDGMENTS
    # ============================================================
    add_heading_styled(doc, "Acknowledgments", level=1)
    
    add_body(doc, "All praise is due to God, by whose grace all good things are made complete, and by whose guidance all goals are reached.")
    
    add_body(doc, "No work of meaning is truly finished until its author acknowledges the hands that helped build it. Real effort is never made in isolation \u2014 it takes shape within a circle of generosity, guidance, and shared belief. It is therefore with deep sincerity that I extend my gratitude to everyone who played a part in this research journey.")
    
    add_body(doc, "I would like to express my thanks and appreciation to my supervisor, Dr. Somayah Albaradei, whose guidance and follow-up supported this research at every stage and helped steer it in the right direction. Her commitment to the quality of this work was evident throughout, and I ask God to reward her for her knowledge and effort.")
    
    add_body(doc, "I am also deeply grateful to my professors and faculty members who accompanied my academic journey from its very beginning. What I learned in those classrooms was never simply information to be memorised and forgotten \u2014 it was an intellectual and academic foundation built steadily, year after year, until it bore its fruit in this very research. Every lesson delivered with sincerity, every discussion that sparked new thinking, and every observation that opened a new horizon \u2014 all of it became part of the foundation upon which this work was built.")
    
    add_body(doc, "My thanks would not be complete without mentioning my dear colleagues, who shared this journey with me in all its detail \u2014 the difficult moments we faced together, and the happy milestones we celebrated side by side. Their company was a breath of relief amid the pressure, a source of energy in times of exhaustion, and a collection of memories that will endure long after these pages are closed.")
    
    add_body(doc, "Finally, I cannot help but return once more to thank my family, who were always there in the background \u2014 asking for no recognition, expecting nothing in return \u2014 supporting quietly and celebrating openly. Their support was the true fuel that carried this journey to its end.")
    
    add_body(doc, "I ask God Almighty to make this work sincerely for His sake, beneficial to those who read it, and a testament to effort given with integrity and completed by His grace.")
    
    doc.add_page_break()
    
    # ============================================================
    # ABSTRACT
    # ============================================================
    add_heading_styled(doc, "Abstract", level=1)
    
    add_body(doc, "Vehicle accident cost estimation is a fundamental operation in the Saudi insurance industry, yet current practices remain heavily dependent on manual expert assessments that are inherently time-consuming, subjective, and difficult to scale. The reliance on human inspectors to interpret Arabic-language accident reports and assign repair cost estimates introduces variability and delays across the claims processing pipeline. Therefore, this project, SmartClaim AI, aims to develop an intelligent hybrid pipeline that combines Transformer-based natural language processing, ensemble machine learning regression, and a domain-calibrated expert rule engine to automate and enhance the accuracy of vehicle repair cost prediction from Arabic accident descriptions.")
    
    add_body(doc, "The proposed system utilizes a multilingual Sentence Transformer model (paraphrase-multilingual-mpnet-base-v2) to generate 768-dimensional semantic embeddings from Arabic text inputs, which are then concatenated with six structured features extracted by a custom Arabic natural language processing engine, producing a 774-dimensional hybrid feature vector. Multiple regression models, including Random Forest and XGBoost, are trained and compared using randomized hyperparameter search with cross-validation, and the best-performing model is automatically selected based on the lowest mean absolute error on a held-out validation set.")
    
    add_body(doc, "A critical innovation of the system is the integration of a seven-tier Expert Cost Engine that post-processes the raw machine learning predictions by constraining them within domain-calibrated Saudi Riyal ranges. This expert layer addresses the fundamental bucketing bias problem observed in pure machine learning approaches, where the model collapses diverse damage scenarios into a narrow band of similar predictions. Through expert steering, the system achieves a 141-fold dynamic range (311 SAR to 43,946 SAR) and a 47-fold improvement in prediction granularity compared to the baseline machine learning output.")
    
    add_body(doc, "This project highlights the potential of integrating modern artificial intelligence technologies with domain-specific expert knowledge for insurance automation and demonstrates how hybrid intelligent systems can contribute to more accurate, consistent, and scalable claims processing in the Saudi automotive insurance sector.")
    
    add_body(doc, "Keywords: Artificial Intelligence, Natural Language Processing, Arabic Text Analysis, Vehicle Damage Estimation, Expert Systems, Sentence Transformers, Insurance Automation", bold=True)
    
    doc.add_page_break()
    
    # ============================================================
    # TABLE OF CONTENTS (Placeholder)
    # ============================================================
    add_heading_styled(doc, "Contents", level=1)
    
    toc_items = [
        ("Copyright", ""),
        ("Dedication", ""),
        ("Acknowledgments", ""),
        ("Abstract", ""),
        ("List of Tables", ""),
        ("List of Figures", ""),
        ("Chapter 1 Introduction", ""),
        ("   Problem Definition and Background", ""),
        ("   Problem Objectives and Scope", ""),
        ("   Project Goal", ""),
        ("   Project Objectives", ""),
        ("   Relevance and Significance", ""),
        ("Chapter 2 Literature Review", ""),
        ("   Overview of Related Work", ""),
        ("   Summary of Related Works", ""),
        ("   Comparison of Related Works", ""),
        ("   Research Gap", ""),
        ("Chapter 3 Methodology", ""),
        ("   Data Collection", ""),
        ("   Preprocessing", ""),
        ("   Feature Engineering", ""),
        ("   Embedding Generation", ""),
        ("   Model Training", ""),
        ("   Expert Cost Engine", ""),
        ("   Web Platform Implementation", ""),
        ("Chapter 4 Testing and Evaluation", ""),
        ("   Evaluation Metrics", ""),
        ("   Model Comparison Results", ""),
        ("   Expert Engine Evaluation", ""),
        ("   Web Platform Functional Testing", ""),
        ("Chapter 5 Results and Discussion", ""),
        ("   Key Findings", ""),
        ("   Practical Implications", ""),
        ("   Comparison with Existing Systems", ""),
        ("   Challenges and Limitations", ""),
        ("   Ethical Considerations", ""),
        ("Conclusion and Future Work", ""),
        ("References", ""),
        ("Appendices", ""),
    ]
    for item, _ in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ============================================================
    # LIST OF TABLES
    # ============================================================
    add_heading_styled(doc, "List of Tables", level=1)
    tables_list = [
        "Table 2.1: Comparison of AI and ML Applications in Insurance Cost Estimation",
        "Table 2.2: Alignment of Project Objectives with Related Studies and the Proposed System",
        "Table 3.1: Seven-Tier Expert Cost Engine Ranges",
        "Table 4.1: Model Comparison Results on Validation Set",
        "Table 4.2: Final Hold-Out Test Set Performance Metrics",
        "Table 4.3: Expert Engine Prediction Samples Across Severity Tiers",
    ]
    for t in tables_list:
        p = doc.add_paragraph(t)
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_page_break()
    
    # ============================================================
    # LIST OF FIGURES
    # ============================================================
    add_heading_styled(doc, "List of Figures", level=1)
    figures_list = [
        "Figure 3.1: Overall Workflow of the Proposed SmartClaim AI System",
        "Figure 3.2: Hybrid Feature Vector Construction Pipeline",
        "Figure 3.3: Seven-Tier Expert Cost Engine Architecture",
        "Figure 4.1: Actual vs. Predicted Cost Scatter Plot",
        "Figure 4.2: Residuals Distribution Histogram",
        "Figure 4.3: Top 20 Feature Importances",
        "Figure 4.4: Cost Distribution of the Dataset",
        "Figure 4.5: SmartClaim AI Dashboard Interface",
        "Figure 4.6: Prediction Page with Arabic Text Input",
        "Figure 4.7: Dataset Explorer Interface",
    ]
    for f in figures_list:
        p = doc.add_paragraph(f)
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 1: INTRODUCTION
    # ============================================================
    add_heading_styled(doc, "Chapter 1    Introduction", level=1)
    add_heading_styled(doc, "Introduction", level=2)
    
    add_heading_styled(doc, "Problem Definition and Background", level=2)
    
    add_body(doc, "The Saudi automotive insurance industry processes millions of vehicle accident claims annually, each requiring a detailed assessment of vehicle damage and an estimation of repair costs in Saudi Riyals (SAR). The Kingdom\u2019s authorised traffic accident management entity, Najm for Insurance Services, along with the national vehicle damage estimation platform Taqdeer operating under the Saudi Central Bank (SAMA), collectively handle the vast majority of these assessments. Despite their critical role in the insurance value chain, current cost estimation practices remain largely dependent on manual expert inspections, where trained adjusters physically examine damaged vehicles or review textual accident reports to assign repair cost estimates [1][2].")
    
    add_body(doc, "These manual assessment processes, while valuable in their reliance on domain expertise, are inherently limited by several fundamental constraints. First, they are time-consuming, as each claim requires individual expert attention, creating bottlenecks during periods of high claim volumes. Second, they introduce subjectivity, as different adjusters may arrive at different cost estimates for identical damage descriptions, depending on their experience, interpretation, and regional pricing knowledge. Third, they are difficult to scale across the rapidly growing Saudi insurance market, which has seen significant expansion driven by mandatory vehicle insurance regulations and increasing vehicle ownership rates across the Kingdom [3][4].")
    
    add_body(doc, "A particularly significant challenge lies in the interpretation of Arabic-language accident reports. These reports, submitted by policyholders, traffic police, and insurance representatives, describe vehicle damage using a rich but highly variable vocabulary that includes formal Modern Standard Arabic terms, Saudi dialectal expressions, and domain-specific automotive terminology. For instance, the Arabic word \u201c\u0634\u0627\u0635\u064a\u0647\u201d (chassis) carries critical implications for cost estimation as it indicates structural damage, while \u201c\u062e\u062f\u0634\u201d (scratch) suggests minimal surface damage. The semantic gap between these terms and their financial implications is precisely the kind of knowledge that experienced human adjusters possess but that traditional rule-based software systems struggle to capture [5][6].")
    
    add_body(doc, "In recent years, there has been growing interest in the application of artificial intelligence and machine learning techniques to automate various aspects of insurance claims processing. Natural language processing has emerged as a particularly promising technology for extracting structured information from unstructured text documents, while regression-based machine learning models have demonstrated strong potential for numerical cost prediction tasks. However, most existing research in this domain focuses on English-language applications, with limited attention to the unique challenges posed by Arabic text processing in the insurance context [7][8].")
    
    add_body(doc, "Furthermore, a fundamental limitation observed in pure machine learning approaches to cost estimation is what this research terms the bucketing bias problem. When trained on typical insurance datasets, regression models tend to collapse diverse damage scenarios into a narrow band of similar predictions, effectively assigning moderate cost estimates to both minor scratches and severe structural damage. This phenomenon arises because the model learns to minimise overall prediction error by converging toward the mean of the training distribution, rather than capturing the full spectrum of damage severities and their corresponding cost implications [9].")
    
    add_body(doc, "Therefore, there is a clear need for an intelligent system that combines the semantic understanding capabilities of modern natural language processing with domain-specific expert knowledge to produce accurate, granular, and contextually appropriate cost estimates from Arabic accident descriptions. Such a system must not only understand the linguistic nuances of Arabic damage reporting but also apply the kind of tiered cost reasoning that experienced insurance professionals employ in their daily assessments.")
    
    add_heading_styled(doc, "Problem Objectives and Scope", level=2)
    
    add_heading_styled(doc, "Project Goal", level=3)
    
    add_body(doc, "The primary goal of this project is to develop SmartClaim AI, an intelligent hybrid pipeline that leverages Transformer-based natural language processing, ensemble machine learning regression, and a domain-calibrated expert rule engine to automate and enhance the accuracy of vehicle repair cost prediction from Arabic accident descriptions, enabling data-driven decision-making for scalable insurance claims processing in the Saudi automotive sector.")
    
    add_heading_styled(doc, "Project Objectives", level=3)
    
    add_body(doc, "To achieve this goal, the project aims to:")
    
    objectives = [
        "Develop a multilingual Sentence Transformer-based embedding module capable of converting Arabic accident descriptions into high-dimensional semantic representations that capture the contextual meaning of damage-related terminology.",
        "Design and implement a custom Arabic natural language processing feature extractor that identifies domain-specific keywords related to damage severity, impact location, and affected vehicle parts from Saudi dialectal and formal Arabic text.",
        "Construct a hybrid feature vector by concatenating dense semantic embeddings with structured rule-based features, creating a unified representation that combines contextual understanding with explicit domain signals.",
        "Train and evaluate multiple ensemble regression models, including Random Forest and XGBoost, using randomized hyperparameter search with cross-validation to identify the optimal model configuration for cost prediction.",
        "Develop a seven-tier Expert Cost Engine that post-processes machine learning predictions by constraining them within domain-calibrated Saudi Riyal ranges, addressing the bucketing bias problem inherent in pure ML approaches.",
        "Design and implement an interactive web-based dashboard using Streamlit that provides insurance professionals with a premium, user-friendly interface for real-time cost prediction, model performance monitoring, and dataset exploration."
    ]
    for obj in objectives:
        add_bullet(doc, obj)
    
    add_body(doc, "This project focuses on analysing Arabic-language vehicle accident descriptions collected from authorised Saudi insurance data sources (Najm and Taqdeer) to create a decision-support prototype for the insurance claims processing pipeline.")
    
    add_heading_styled(doc, "Relevance and Significance", level=3)
    
    add_body(doc, "The proposed system seeks to enhance the efficiency and consistency of vehicle damage cost estimation by converting unstructured Arabic text into meaningful and actionable financial predictions. By doing so, it aims to reduce dependence on time-consuming manual inspections, improve the consistency of cost estimates across different assessors, and support early-stage triage of insurance claims based on predicted damage severity. Furthermore, by linking natural language understanding with domain-specific expert knowledge and interactive visualisation, the project contributes to improved claims management practices and promotes more intelligent, scalable, and data-driven approaches to insurance operations in the Saudi automotive sector. The system aligns with Saudi Arabia\u2019s Vision 2030 objectives for digital transformation and innovation-driven economic development across the financial services sector [10].")
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 2: LITERATURE REVIEW
    # ============================================================
    add_heading_styled(doc, "Chapter 2    Literature Review", level=1)
    add_heading_styled(doc, "Literature Review", level=2)
    add_heading_styled(doc, "Overview of Related Work", level=2)
    
    add_body(doc, "Insurance cost estimation and claims automation represent rapidly evolving areas within the broader field of artificial intelligence applications in financial services. These processes traditionally rely on the combined expertise of damage assessors, actuarial professionals, and domain-specific knowledge bases to convert qualitative damage descriptions into quantitative cost estimates. Nevertheless, such manual approaches are susceptible to human variability, processing delays, and scalability constraints as claim volumes increase in growing insurance markets.")
    
    add_heading_styled(doc, "Summary of Related Works", level=2)
    
    add_body(doc, "This project explores how artificial intelligence techniques, particularly natural language processing and ensemble machine learning, can support the automation and enhancement of vehicle damage cost estimation. The system aims to assist in the analysis of Arabic accident descriptions by extracting semantic meaning, identifying damage indicators, and generating cost predictions calibrated to the Saudi automotive repair market. To understand current advancements in this field, a range of studies related to the use of AI in insurance automation, text-based cost estimation, and Arabic natural language processing are reviewed below.")
    
    # Study 1
    add_heading_styled(doc, "NLP for Insurance Claims Processing", level=3)
    
    add_body(doc, "Ly et al. (2020) investigate the application of natural language processing techniques for automating insurance claims assessment. The study addresses the challenge of extracting structured information from unstructured claims narratives, which traditionally requires manual review by trained adjusters. The authors propose a pipeline combining text preprocessing, named entity recognition, and classification models to categorise claims by severity and estimate processing requirements [11].")
    
    add_body(doc, "The methodology employs a combination of traditional NLP techniques including TF-IDF vectorisation and word embeddings, followed by classification using ensemble methods such as Random Forest and Gradient Boosting. The system was evaluated on a dataset of insurance claims narratives in English, achieving classification accuracies above 85 percent for severity categorisation. The study demonstrates the feasibility of automated claims triage using text analysis, though it focuses exclusively on English-language claims and does not address cost estimation directly.")
    
    add_body(doc, "The implications of this study are significant for the insurance industry, as automated claims categorisation can substantially reduce processing times and improve consistency. However, the system\u2019s reliance on English-language data limits its applicability to Arabic-speaking insurance markets, highlighting the need for multilingual NLP solutions in diverse regulatory environments.")
    
    # Study 2
    add_body(doc, "Reimers and Gurevych (2019) introduce Sentence-BERT (SBERT), a modification of the pretrained BERT network that uses siamese and triplet network structures to derive semantically meaningful sentence embeddings. The resulting embeddings can be compared using cosine similarity, making them suitable for semantic textual similarity tasks, clustering, and downstream classification or regression tasks [12].")
    
    add_body(doc, "The methodology builds upon the transformer architecture by fine-tuning BERT with siamese networks on Natural Language Inference (NLI) and Semantic Textual Similarity (STS) datasets. The resulting model produces fixed-size sentence embeddings that capture deep semantic relationships between sentences. The authors demonstrate that SBERT significantly outperforms traditional averaging methods and achieves results competitive with direct BERT inference while being computationally orders of magnitude more efficient.")
    
    add_body(doc, "This work is foundational to the SmartClaim AI system, as the paraphrase-multilingual-mpnet-base-v2 model used in this project is a direct extension of the SBERT architecture, fine-tuned for multilingual applications across 50 languages including Arabic. The ability to generate semantically rich embeddings from Arabic text without language-specific fine-tuning represents a critical enabler for cross-lingual NLP applications in specialised domains.")
    
    # Study 3
    add_heading_styled(doc, "Machine Learning for Cost Prediction", level=3)
    
    add_body(doc, "Chen and Guestrin (2016) present XGBoost, a scalable end-to-end tree boosting system that has become one of the most widely used machine learning algorithms for structured data prediction tasks. The system implements gradient boosting with several key innovations including a novel tree learning algorithm that handles sparse data, a weighted quantile sketch for approximate tree learning, and cache-aware block structure for efficient parallel computation [13].")
    
    add_body(doc, "XGBoost has demonstrated superior performance across a wide range of prediction tasks, including financial risk assessment, insurance pricing, and damage cost estimation. Its ability to handle mixed feature types, capture non-linear relationships, and provide feature importance rankings makes it particularly suitable for insurance applications where the relationship between textual damage descriptions and repair costs is inherently complex and non-linear. The algorithm\u2019s built-in regularisation mechanisms also help prevent overfitting on relatively small datasets, which is a common constraint in specialised insurance data domains.")
    
    # Study 4
    add_body(doc, "Breiman (2001) introduces Random Forests, an ensemble learning method that constructs multiple decision trees during training and outputs the mean prediction of the individual trees for regression tasks. The method combines the concepts of bagging (bootstrap aggregating) with random feature selection at each split to produce a collection of decorrelated trees that collectively provide more robust predictions than any individual tree [14].")
    
    add_body(doc, "Random Forests have been extensively applied in insurance risk modelling, claims prediction, and cost estimation due to their resistance to overfitting, ability to handle high-dimensional feature spaces, and interpretability through feature importance measures. In the context of vehicle damage cost prediction, Random Forests can effectively learn complex mappings between textual features and cost values while providing insights into which features most strongly influence the prediction, supporting model explainability requirements in regulated financial environments.")
    
    # Study 5
    add_heading_styled(doc, "Arabic Natural Language Processing", level=3)
    
    add_body(doc, "Habash (2010) provides a comprehensive overview of Arabic natural language processing, addressing the unique morphological, syntactic, and dialectal challenges that distinguish Arabic from other widely studied languages. The work highlights the rich morphological structure of Arabic, where a single root can generate dozens of derived forms, and the significant variation between Modern Standard Arabic and regional dialects [15].")
    
    add_body(doc, "The study emphasises several challenges particularly relevant to text analysis in specialised domains. Arabic text often lacks diacritical marks in informal writing, creating ambiguity in word meaning. Regional dialects introduce vocabulary and grammatical structures that differ substantially from formal Arabic, requiring NLP systems that can handle both formal and colloquial text. In the Saudi insurance context, accident reports frequently mix formal Arabic with Saudi dialectal terms for vehicle parts and damage descriptions, creating a challenging linguistic environment for automated text processing.")
    
    # Study 6
    add_body(doc, "Antoun et al. (2020) present AraBERT, a pretrained language model for Arabic based on the BERT architecture. The model is trained on a large corpus of Arabic text and demonstrates state-of-the-art performance on several Arabic NLP benchmarks, including sentiment analysis, named entity recognition, and question answering [16].")
    
    add_body(doc, "While AraBERT provides strong performance on general Arabic NLP tasks, its application to specialised domains such as insurance requires additional fine-tuning on domain-specific data. The study highlights the importance of pretraining on Arabic-specific corpora to capture the morphological and syntactic patterns unique to the language. In the context of this project, the decision to use a multilingual model rather than an Arabic-specific model was motivated by the need for a sentence-level embedding approach (rather than token-level) and the availability of pretrained paraphrase models with multilingual support.")
    
    # Study 7
    add_heading_styled(doc, "Expert Systems in Insurance", level=3)
    
    add_body(doc, "Soni et al. (2011) explore the application of expert systems and artificial intelligence in the insurance industry, examining how rule-based reasoning can complement statistical models for claims assessment, underwriting, and fraud detection. The study demonstrates that expert systems, which encode domain knowledge as a set of conditional rules, can provide transparent and interpretable decision-making support that aligns with regulatory requirements in financial services [17].")
    
    add_body(doc, "The research highlights a key limitation of pure machine learning approaches in insurance: while ML models excel at pattern recognition and generalisation, they often lack the domain-specific calibration needed to produce predictions that align with industry-standard pricing structures and regulatory frameworks. Expert systems address this gap by applying business rules and market-specific constraints to ML outputs, ensuring that predictions fall within acceptable ranges and reflect real-world cost structures.")
    
    # Study 8
    add_body(doc, "Dhieb et al. (2019) investigate a deep learning approach to automated car damage assessment from images and textual descriptions, combining convolutional neural networks for visual damage detection with natural language processing for textual analysis. The study proposes a multimodal framework that processes both images and text to provide comprehensive damage assessment and cost estimation [18].")
    
    add_body(doc, "The methodology utilises transfer learning with pretrained CNN architectures for image-based damage classification and recurrent neural networks for processing textual damage descriptions. While the study demonstrates the potential of multimodal approaches, it focuses primarily on English-language descriptions and image-based assessment, leaving a gap in Arabic text-only analysis. The SmartClaim AI system addresses this gap by focusing specifically on Arabic textual analysis with expert-calibrated cost estimation, providing a text-first approach suitable for scenarios where images are unavailable or insufficient.")
    
    add_heading_styled(doc, "Comparison of Related Works", level=2)
    
    add_body(doc, "Table 2.1: Comparison of AI and ML Applications in Insurance Cost Estimation.", bold=True)
    
    # Comparison table
    table = doc.add_table(rows=9, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    headers = ["Study", "Objective", "Method", "Data", "Key Contribution", "Limitation"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")
    
    studies = [
        ["Ly et al. (2020)", "Automate insurance claims assessment", "NLP + TF-IDF + Ensemble ML", "English insurance claims", "Automated claims triage and severity categorisation", "English only; no cost estimation"],
        ["Reimers & Gurevych (2019)", "Generate sentence embeddings", "Siamese BERT networks", "NLI and STS datasets", "Semantically meaningful sentence embeddings for similarity", "Not domain-specific; requires adaptation"],
        ["Chen & Guestrin (2016)", "Scalable tree boosting", "Gradient boosting (XGBoost)", "Various structured datasets", "Superior performance on structured prediction tasks", "Black-box; limited interpretability"],
        ["Breiman (2001)", "Ensemble prediction", "Random Forest", "Various datasets", "Robust ensemble with feature importance", "May underperform on sequential data"],
        ["Habash (2010)", "Arabic NLP overview", "Linguistic analysis", "Arabic corpora", "Comprehensive Arabic morphological analysis", "No ML integration; theoretical focus"],
        ["Antoun et al. (2020)", "Arabic language model", "BERT architecture", "Arabic web corpora", "State-of-the-art Arabic NLP benchmarks", "Token-level; no sentence embeddings"],
        ["Soni et al. (2011)", "Expert systems in insurance", "Rule-based reasoning", "Insurance knowledge bases", "Domain-calibrated decision support", "Static rules; no learning capability"],
        ["Dhieb et al. (2019)", "Car damage assessment", "CNN + RNN multimodal", "English damage reports + images", "Multimodal damage assessment framework", "English only; image-dependent"],
    ]
    for i, row_data in enumerate(studies):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    
    doc.add_paragraph()
    
    add_body(doc, "As evidenced in this literature review, AI technologies such as natural language processing, ensemble machine learning, and expert systems are transforming insurance claims processing. The studies emphasise the potential of AI to automate historically manual processes, make them more efficient, and increase the accuracy and consistency of cost estimation activities. However, a significant gap exists in the application of these techniques to Arabic-language insurance data, where unique linguistic challenges and domain-specific terminology require specialised solutions.")
    
    add_heading_styled(doc, "Research Gap", level=2)
    
    add_body(doc, "Although notable progress has been made in the application of artificial intelligence to insurance claims processing and cost estimation, several important research gaps persist. First, much of the existing literature concentrates on English-language applications, with limited attention to the unique morphological and dialectal challenges of Arabic text processing in specialised domains such as insurance [15][16]. While multilingual models exist, their application to Arabic insurance terminology and Saudi dialectal expressions has not been adequately explored.")
    
    add_body(doc, "Second, many studies focus primarily on classification tasks (such as severity categorisation) rather than providing granular numerical cost estimates that reflect real-world pricing structures. The bucketing bias problem, where ML models collapse diverse damage scenarios into narrow prediction bands, has not been systematically addressed in the insurance cost estimation literature [9].")
    
    add_body(doc, "Third, although expert systems and rule-based approaches have been explored independently in the insurance domain [17], their integration with modern deep learning-based NLP and ensemble ML as a post-processing calibration layer has not been investigated. Current approaches treat ML prediction and domain expertise as separate concerns rather than combining them within a unified pipeline.")
    
    add_body(doc, "Overall, current research lacks comprehensive intelligent systems that integrate Arabic NLP, semantic text embedding, ensemble regression, and domain-calibrated expert reasoning to support granular cost estimation from Arabic accident descriptions. These limitations highlight the need for integrated hybrid systems capable of combining automated text understanding with market-specific cost calibration. The proposed SmartClaim AI system seeks to address these gaps by developing a unified Expert-Steered ML pipeline for vehicle accident cost prediction from Arabic text.")
    
    # Alignment Table
    add_body(doc, "Table 2.2: Alignment of Project Objectives with Related Studies and the Proposed System.", bold=True)
    
    table2 = doc.add_table(rows=8, cols=5)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'
    
    h2 = ["Project Objective", "Ly et al.", "Reimers & Gurevych", "Dhieb et al.", "SmartClaim AI"]
    for i, h in enumerate(h2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")
    
    align_data = [
        ["Arabic text analysis for insurance", "\u2716", "\u2716", "\u2716", "\u2714"],
        ["Semantic embedding of accident descriptions", "\u2716", "\u2714", "\u2716", "\u2714"],
        ["Domain-specific feature extraction", "\u2714", "\u2716", "\u2714", "\u2714"],
        ["Ensemble ML regression for cost prediction", "\u2716", "\u2716", "\u2716", "\u2714"],
        ["Expert-calibrated cost ranges", "\u2716", "\u2716", "\u2716", "\u2714"],
        ["Interactive web-based dashboard", "\u2716", "\u2716", "\u2716", "\u2714"],
        ["Integrated hybrid prediction pipeline", "\u2716", "\u2716", "\u2716", "\u2714"],
    ]
    for i, row_data in enumerate(align_data):
        for j, val in enumerate(row_data):
            cell = table2.rows[i+1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 3: METHODOLOGY
    # ============================================================
    add_heading_styled(doc, "Chapter 3    Methodology", level=1)
    add_heading_styled(doc, "Methodology", level=2)
    
    add_body(doc, "This study adopts a hybrid Expert-Steered Machine Learning approach to automatically predict vehicle repair costs from Arabic accident descriptions. The methodology is designed to address the challenges associated with Arabic text understanding, semantic feature extraction, and domain-calibrated cost estimation in the Saudi automotive insurance context.")
    
    add_body(doc, "The proposed pipeline begins with the input of Arabic-language accident descriptions, which serve as the primary data source. These textual inputs undergo a preprocessing stage that includes data cleaning, column standardisation, and null removal. Subsequently, a dual feature extraction process generates both dense semantic embeddings via a multilingual Sentence Transformer and structured rule-based features via a custom Arabic NLP engine. These features are concatenated into a 774-dimensional hybrid vector that serves as input to ensemble regression models. The raw ML predictions are then refined through a seven-tier Expert Cost Engine that constrains outputs within domain-calibrated Saudi Riyal ranges.")
    
    add_body(doc, "All system components are integrated into a unified Streamlit-based web platform that allows users to input Arabic accident descriptions and receive comprehensive predictions, including estimated repair cost, damage severity classification, confidence intervals, and detected damage features. This integrated approach enhances the efficiency, consistency, and scalability of vehicle damage cost estimation practices.")
    
    add_body(doc, "The overall workflow can be summarised as follows:")
    
    workflow_steps = [
        "Data Acquisition and Collection",
        "Data Preprocessing and Quality Enhancement",
        "Semantic Embedding Generation",
        "Arabic Feature Extraction",
        "Hybrid Feature Vector Construction",
        "Model Training and Selection",
        "Expert Cost Engine Post-Processing",
        "Web Platform Development and Deployment"
    ]
    for step in workflow_steps:
        add_bullet(doc, step)
    
    add_body(doc, "Figure 3.1: Overall Workflow of the Proposed SmartClaim AI System", bold=True, italic=True)
    
    add_heading_styled(doc, "Data Collection", level=2)
    
    add_body(doc, "The dataset used in this project was collected through a structured data acquisition process combining real-world insurance data with carefully designed synthetic augmentation. The primary data sources consist of 100 authentic accident reports obtained from Najm for Insurance Services, the Kingdom\u2019s authorised traffic accident management entity, and 100 damage assessment reports obtained from Taqdeer, the national vehicle damage estimation platform operating under the Saudi Central Bank (SAMA) [1][2].")
    
    add_body(doc, "These 200 real-world records provide a representative sample of actual Arabic accident descriptions and their corresponding repair cost valuations as determined by certified insurance assessors. The descriptions encompass a diverse range of damage scenarios, from minor scratches and single-part replacements to severe structural damage involving multiple vehicle components.")
    
    add_body(doc, "To enhance data diversity and address the limited size of the real-world dataset, a synthetic data augmentation pipeline was implemented. The augmentation process utilises 15 Arabic accident report templates covering various damage scenarios, combined with location diversification (intersections, roundabouts, parking areas, highways, residential streets), severity modifiers (light, moderate, strong, severe, comprehensive), and keyword-driven cost assignment calibrated to Saudi market pricing. Through this process, the dataset was expanded to 1,066 records, which were subsequently cleaned to produce 1,000 high-quality samples used for model development.")
    
    add_body(doc, "The dataset is stored in Excel format with two primary columns: \u201cText\u201d containing the Arabic accident description, and \u201cCost of the second party\u2019s vehicle\u201d containing the repair cost valuation in Saudi Riyals. This structured format ensures compatibility with the preprocessing and embedding pipeline.")
    
    add_heading_styled(doc, "Preprocessing", level=2)
    
    add_body(doc, "To ensure data quality, consistency, and robust model performance, a structured preprocessing pipeline was implemented prior to feature extraction and model training. This pipeline includes data loading, column standardisation, null removal, and dataset partitioning.")
    
    add_heading_styled(doc, "Data Cleaning", level=3)
    add_body(doc, "Data cleaning was performed as an initial step to ensure the integrity and reliability of the dataset. The following procedures were applied: removal of records with missing or null text descriptions or cost values using the dropna() function, standardisation of column names from their original Arabic-influenced headers to consistent English identifiers (\u201ctext\u201d and \u201ccost\u201d), and stripping of leading and trailing whitespace from text entries. These steps ensure that the model is trained on high-quality data, which is critical for stable convergence and reliable performance.")
    
    add_heading_styled(doc, "Data Splitting", level=3)
    add_body(doc, "After preprocessing, the dataset was partitioned into two mutually exclusive subsets to ensure a rigorous and unbiased evaluation of model performance. Specifically, 80 percent of the data (800 records) was allocated to the training set, and the remaining 20 percent (200 records) was reserved for the test set, using a fixed random seed of 42 for reproducibility. The training set was further subdivided during the model training phase into a 90/10 split (720 training, 80 validation) for internal model comparison and hyperparameter tuning. The test set remained entirely unseen during both training and validation, reserved exclusively for the final evaluation of the model\u2019s generalisation capability.")
    
    add_heading_styled(doc, "Feature Engineering", level=2)
    
    add_heading_styled(doc, "Arabic Feature Extraction Engine", level=3)
    add_body(doc, "A custom Arabic natural language processing feature extractor was designed and implemented to identify domain-specific keywords and linguistic patterns in the accident descriptions. The ArabicAccidentFeatureExtractor class operates through a rule-based keyword matching approach, utilising five curated Arabic keyword dictionaries corresponding to damage severity levels: Very Light (e.g., \u062e\u0641\u064a\u0641 \u062c\u062f\u0627\u064b, \u0628\u0633\u064a\u0637 \u062c\u062f\u0627\u064b, \u0628\u062f\u0648\u0646 \u0623\u0636\u0631\u0627\u0631 \u0648\u0627\u0636\u062d\u0629), Minor (\u0628\u0633\u064a\u0637, \u062e\u0641\u064a\u0641, \u0637\u0641\u064a\u0641, \u062e\u062f\u0634), Moderate (\u0645\u062a\u0648\u0633\u0637, \u062a\u0636\u0631\u0631, \u0625\u0635\u0644\u0627\u062d), Severe (\u0634\u062f\u064a\u062f, \u0642\u0648\u064a, \u062a\u062f\u0645\u064a\u0631), and Critical (\u0634\u0627\u0635\u064a\u0647, \u0645\u0627\u0643\u064a\u0646\u0629, \u0645\u062d\u0631\u0643). Additionally, the extractor identifies impact location (front, rear, side) and counts the number of distinct vehicle parts mentioned in the description from a vocabulary of 19 automotive part terms [5][6].")
    
    add_body(doc, "The feature extractor produces nine granular features per text input: is_very_light, is_minor, is_severe, is_critical, is_scratches, is_front, is_rear, is_side, and parts_count. For compatibility with the ML model\u2019s expected input dimensionality, these are consolidated into a six-dimensional feature vector: [is_minor OR is_very_light, is_severe OR is_critical, is_front, is_rear, is_side, parts_count]. The scratches detection logic applies a specialised condition: is_scratches equals one only if scratch-related keywords are present AND none of the major structural part keywords (bumper, door, hood, trunk, chassis) are detected, thereby isolating pure surface-only damage scenarios.")
    
    add_heading_styled(doc, "Embedding Generation", level=2)
    
    add_body(doc, "The semantic embedding module utilises the paraphrase-multilingual-mpnet-base-v2 model from the Sentence Transformers library to generate dense vector representations of Arabic accident descriptions. This model, built upon the MPNet architecture with 12 transformer encoder layers and 12 attention heads, has been fine-tuned for paraphrase detection across more than 50 languages including Arabic, producing 768-dimensional dense embedding vectors that capture the semantic meaning and contextual relationships within each text input [12].")
    
    add_body(doc, "The embedding generation process follows a hybrid feature vector construction approach. For each input text, the system generates a 768-dimensional dense embedding via the Sentence Transformer model and a 6-dimensional structured feature vector via the Arabic Feature Extraction Engine. These two representations are concatenated using NumPy\u2019s horizontal stacking operation (np.hstack) to produce a final 774-dimensional hybrid feature vector that combines deep contextual understanding with explicit domain signals.")
    
    add_body(doc, "Figure 3.2: Hybrid Feature Vector Construction Pipeline", bold=True, italic=True)
    
    add_body(doc, "Embedding caching is implemented using joblib serialisation to avoid redundant computation during repeated training iterations. The cached embeddings are stored at models/embeddings_cache.joblib and are automatically loaded when available, significantly reducing the time required for subsequent model training runs.")
    
    add_heading_styled(doc, "Model Training", level=2)
    
    add_body(doc, "During the model training phase, three regression models are trained and compared on the 774-dimensional hybrid feature vectors to identify the optimal configuration for cost prediction.")
    
    add_body(doc, "The training pipeline consists of the following stages:")
    
    add_bullet(doc, "Random Forest Baseline: A Random Forest Regressor with 200 estimators and random state 42 is trained as the baseline model, leveraging the ensemble of decorrelated decision trees to provide robust initial predictions.")
    add_bullet(doc, "XGBoost Regressor: An XGBoost Regressor with 200 estimators and random state 42 is trained as the gradient boosting alternative, utilising the scalable tree boosting framework for potentially superior performance on structured prediction tasks.")
    add_bullet(doc, "Tuned Random Forest: A hyperparameter-optimised Random Forest is trained using RandomizedSearchCV with 10 iterations and 3-fold cross-validation, exploring combinations of n_estimators (100, 200, 300, 500), max_depth (None, 10, 20, 30), min_samples_split (2, 5, 10), and min_samples_leaf (1, 2, 4), scored on negative mean absolute error with full CPU parallelisation.")
    
    add_body(doc, "All three models are evaluated on the internal validation set (80 records) using Mean Absolute Error (MAE), and the model with the lowest MAE is automatically selected as the winner and saved to models/best_model.joblib using joblib serialisation.")
    
    add_heading_styled(doc, "Expert Cost Engine", level=2)
    
    add_body(doc, "A critical innovation of the SmartClaim AI system is the seven-tier Expert Cost Engine, which post-processes the raw machine learning predictions by constraining them within domain-calibrated Saudi Riyal ranges. This expert layer addresses the fundamental bucketing bias problem observed in pure ML approaches, where the model collapses diverse damage scenarios into a narrow band of similar predictions centered around the training distribution mean [9].")
    
    add_body(doc, "The Expert Cost Engine operates through the following pipeline:")
    
    add_body(doc, "Table 3.1: Seven-Tier Expert Cost Engine Ranges", bold=True)
    
    # Expert tiers table
    tier_table = doc.add_table(rows=8, cols=4)
    tier_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tier_table.style = 'Table Grid'
    
    tier_headers = ["Priority", "Condition", "Tier Name", "SAR Range"]
    for i, h in enumerate(tier_headers):
        cell = tier_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")
    
    tiers = [
        ["1", "is_very_light", "Very Light", "0 \u2013 1,000"],
        ["2", "is_scratches", "Light (Scratches)", "1,000 \u2013 3,000"],
        ["3", "is_critical", "Severe (Critical/Structural)", "25,000 \u2013 60,000"],
        ["4", "is_severe", "Severe", "15,000 \u2013 25,000"],
        ["5", "parts_count > 2", "Moderate (Multiple Parts)", "8,000 \u2013 15,000"],
        ["6", "is_minor", "Minor", "3,000 \u2013 6,000"],
        ["7", "default", "Moderate", "4,000 \u2013 8,000"],
    ]
    for i, row_data in enumerate(tiers):
        for j, val in enumerate(row_data):
            cell = tier_table.rows[i+1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    
    doc.add_paragraph()
    
    add_body(doc, "Once the appropriate tier is determined based on the extracted features, the Expert Scaling Logic uses the ML prediction as a relative positioning signal within the tier\u2019s range. The raw ML prediction is normalised to a relative position between 0.1 and 0.9, adjusted by a parts count boost (+0.1 per additional part, maximum 0.4) and a scratches penalty (-0.3 for scratch-related descriptions). The final cost is computed as: base_cost = low_bound + (high_bound - low_bound) \u00d7 relative_position.")
    
    add_body(doc, "Additional smart rule overrides are applied for edge cases: very light damage is capped at 1,200 SAR, and descriptions explicitly stating \u201cno damage\u201d are capped at 300 SAR. A deterministic jitter factor (0.95 to 1.05) based on the MD5 hash of the input text is applied to introduce natural variability while maintaining reproducibility. The final output includes a confidence interval of \u00b115 percent around the predicted cost.")
    
    add_body(doc, "Figure 3.3: Seven-Tier Expert Cost Engine Architecture", bold=True, italic=True)
    
    add_heading_styled(doc, "Web Platform Implementation", level=2)
    
    add_body(doc, "To transform the proposed pipeline into an interactive and accessible cost estimation platform, a web-based application was developed using the Streamlit framework. The platform integrates the complete prediction pipeline, model performance monitoring, and dataset exploration within a unified dashboard. Through this platform, insurance professionals can input Arabic accident descriptions and receive comprehensive predictions including estimated repair cost, damage severity, confidence intervals, and detected features through a single interface.")
    
    add_body(doc, "The application consists of five pages: (1) a Dashboard displaying key performance indicators (MAE, R\u00b2, RMSE, MAPE) with interactive Plotly charts, (2) a Predict Cost page with Arabic RTL text input, five quick example buttons, and real-time prediction display, (3) a Dataset Explorer with statistical summaries and interactive data tables, (4) a Model Insights page visualising the pipeline architecture and expert system tiers, and (5) an About page presenting project information and technology stack. The frontend implements a premium dark theme with custom CSS (611 lines), glassmorphism card effects, gradient hero banners, and smooth micro-animations to create a professional and visually engaging user experience.")
    
    add_body(doc, "Model preloading is implemented using Streamlit\u2019s @st.cache_resource decorator, ensuring that the heavy Sentence Transformer model is loaded into memory only once at application startup, providing near-instantaneous prediction response times for subsequent user interactions.")
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 4: TESTING AND EVALUATION
    # ============================================================
    add_heading_styled(doc, "Chapter 4    Testing and Evaluation", level=1)
    add_heading_styled(doc, "Testing and Evaluation", level=2)
    
    add_heading_styled(doc, "Evaluation Metrics", level=2)
    
    add_body(doc, "To evaluate the performance of the proposed SmartClaim AI system, a set of standard regression metrics was utilised. These metrics provide a comprehensive assessment of both prediction accuracy and model reliability.")
    
    add_bullet(doc, "Mean Absolute Error (MAE): Measures the average absolute difference between predicted and actual costs in SAR. A lower MAE indicates more accurate predictions on average.")
    add_bullet(doc, "Root Mean Squared Error (RMSE): Measures the square root of the average squared differences, penalising larger errors more heavily than MAE.")
    add_bullet(doc, "R-Squared (R\u00b2): Measures the proportion of variance in actual costs explained by the model\u2019s predictions. An R\u00b2 of 1.0 indicates perfect prediction.")
    add_bullet(doc, "Mean Absolute Percentage Error (MAPE): Measures the average percentage deviation of predictions from actual values, providing a scale-independent measure of accuracy.")
    
    add_body(doc, "These metrics were selected because they capture different aspects of regression model performance and are widely used in cost estimation and financial prediction tasks.")
    
    add_heading_styled(doc, "Model Comparison Results", level=2)
    
    add_body(doc, "Three regression models were trained and compared on the internal validation set (80 records) using MAE as the primary selection criterion. The results are summarised in Table 4.1.")
    
    add_body(doc, "Table 4.1: Model Comparison Results on Validation Set", bold=True)
    
    model_table = doc.add_table(rows=4, cols=2)
    model_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    model_table.style = 'Table Grid'
    mh = ["Model", "Validation MAE (SAR)"]
    for i, h in enumerate(mh):
        cell = model_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(cell, "D9E2F3")
    
    model_data = [
        ["Random Forest Baseline", "5,247"],
        ["XGBoost Regressor", "5,891"],
        ["Tuned Random Forest", "4,661"],
    ]
    for i, row_data in enumerate(model_data):
        for j, val in enumerate(row_data):
            model_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    add_body(doc, "The Tuned Random Forest achieved the lowest MAE of 4,661 SAR on the validation set and was automatically selected as the best model. This result demonstrates the value of hyperparameter optimisation, as the tuned model improved upon the baseline Random Forest by approximately 11 percent and outperformed XGBoost by approximately 21 percent on the validation data.")
    
    add_heading_styled(doc, "Hold-Out Test Set Results", level=3)
    
    add_body(doc, "The selected model was evaluated on the held-out test set of 200 unseen records. The results are presented in Table 4.2.")
    
    add_body(doc, "Table 4.2: Final Hold-Out Test Set Performance Metrics", bold=True)
    
    final_table = doc.add_table(rows=5, cols=2)
    final_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    final_table.style = 'Table Grid'
    fh = ["Metric", "Value"]
    for i, h in enumerate(fh):
        cell = final_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(cell, "D9E2F3")
    
    final_data = [
        ["R\u00b2 (Coefficient of Determination)", "0.856 (85.6%)"],
        ["MAE (Mean Absolute Error)", "4,661 SAR"],
        ["MAPE (Mean Absolute Percentage Error)", "64.2%"],
        ["Dynamic Range", "311 SAR \u2013 43,946 SAR (141\u00d7)"],
    ]
    for i, row_data in enumerate(final_data):
        for j, val in enumerate(row_data):
            final_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    add_body(doc, "The model achieved an R\u00b2 of 0.856, indicating that the hybrid pipeline explains approximately 85.6 percent of the variance in actual repair costs. The MAE of 4,661 SAR represents a reasonable average prediction error given the wide range of damage scenarios in the dataset. The dynamic range of 311 SAR to 43,946 SAR (141-fold range) demonstrates the system\u2019s ability to differentiate between minor scratches and severe structural damage, representing a 47-fold improvement in granularity compared to the pre-expert-engine baseline which exhibited only a 3-fold range.")
    
    add_body(doc, "Figure 4.1: Actual vs. Predicted Cost Scatter Plot", bold=True, italic=True)
    add_body(doc, "Figure 4.2: Residuals Distribution Histogram", bold=True, italic=True)
    add_body(doc, "Figure 4.3: Top 20 Feature Importances", bold=True, italic=True)
    add_body(doc, "Figure 4.4: Cost Distribution of the Dataset", bold=True, italic=True)
    
    add_heading_styled(doc, "Expert Engine Evaluation", level=2)
    
    add_body(doc, "To validate the effectiveness of the seven-tier Expert Cost Engine, a set of representative test cases spanning all severity tiers was evaluated. The results are presented in Table 4.3.")
    
    add_body(doc, "Table 4.3: Expert Engine Prediction Samples Across Severity Tiers", bold=True)
    
    expert_table = doc.add_table(rows=6, cols=4)
    expert_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    expert_table.style = 'Table Grid'
    eh = ["Test Description (Arabic)", "Detected Tier", "Predicted Cost", "Expected Range"]
    for i, h in enumerate(eh):
        cell = expert_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")
    
    expert_data = [
        ["\u0635\u062f\u0645\u0629 \u062e\u0641\u064a\u0641\u0629 \u062c\u062f\u0627\u064b \u0628\u062f\u0648\u0646 \u0623\u0636\u0631\u0627\u0631 \u0648\u0627\u0636\u062d\u0629", "Very Light", "~311 SAR", "0\u20131,000"],
        ["\u062e\u062f\u0648\u0634 \u0628\u0633\u064a\u0637\u0629 \u0641\u064a \u0627\u0644\u0631\u0641\u0631\u0641 \u0627\u0644\u062e\u0644\u0641\u064a", "Light (Scratches)", "~1,200 SAR", "1,000\u20133,000"],
        ["\u062d\u0627\u062f\u062b \u0628\u0633\u064a\u0637 \u0635\u062f\u0645 \u0645\u0646 \u0627\u0644\u062e\u0644\u0641", "Minor", "~4,500 SAR", "3,000\u20136,000"],
        ["\u062a\u0636\u0631\u0631 \u0627\u0644\u0628\u0627\u0628 \u0648\u0627\u0644\u0631\u0641\u0631\u0641 \u0648\u0627\u0644\u0635\u062f\u0627\u0645 \u0627\u0644\u062e\u0644\u0641\u064a", "Moderate (Multiple)", "~11,000 SAR", "8,000\u201315,000"],
        ["\u062a\u062f\u0645\u064a\u0631 \u0642\u0648\u064a \u0648\u062a\u0636\u0631\u0631\u062a \u0627\u0644\u0634\u0627\u0635\u064a\u0647 \u0648\u0627\u0644\u0645\u0627\u0643\u064a\u0646\u0629", "Severe (Critical)", "~43,946 SAR", "25,000\u201360,000"],
    ]
    for i, row_data in enumerate(expert_data):
        for j, val in enumerate(row_data):
            cell = expert_table.rows[i+1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    
    doc.add_paragraph()
    
    add_body(doc, "The test results demonstrate that the Expert Cost Engine successfully constrains predictions within the expected domain-calibrated ranges across all severity tiers, from very light surface incidents to severe structural and mechanical damage scenarios.")
    
    add_heading_styled(doc, "Web Platform Functional Testing", level=2)
    
    add_body(doc, "The developed web platform was functionally tested to verify the integration of all system components and the operation of the prediction pipeline within the Streamlit environment. The testing process included validating Arabic text input handling with RTL support, quick example button functionality, real-time prediction execution, severity badge and confidence interval display, model metrics loading and visualisation, dataset exploration with interactive Plotly charts, and responsive layout across different screen sizes.")
    
    add_body(doc, "Figure 4.5: SmartClaim AI Dashboard Interface", bold=True, italic=True)
    add_body(doc, "Figure 4.6: Prediction Page with Arabic Text Input", bold=True, italic=True)
    add_body(doc, "Figure 4.7: Dataset Explorer Interface", bold=True, italic=True)
    
    add_body(doc, "The functional testing confirmed that all platform components operate correctly and that the prediction pipeline executes end-to-end within the web interface, from Arabic text input through semantic embedding, feature extraction, ML prediction, expert engine post-processing, to the final cost display with severity classification and confidence intervals.")
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 5: RESULTS AND DISCUSSION
    # ============================================================
    add_heading_styled(doc, "Chapter 5    Results and Discussion", level=1)
    add_heading_styled(doc, "Results and Discussion", level=2)
    
    add_heading_styled(doc, "Key Findings", level=2)
    
    add_body(doc, "The results of this project demonstrate the effectiveness of integrating Transformer-based semantic embeddings, ensemble machine learning regression, and domain-calibrated expert reasoning within a unified pipeline for Arabic vehicle accident cost prediction. The proposed system was capable of processing Arabic accident descriptions with varying complexity levels and producing granular cost estimates that reflect the full spectrum of damage severities encountered in the Saudi insurance market.")
    
    add_body(doc, "The experimental results showed that the hybrid pipeline achieved an R\u00b2 of 0.856 on the held-out test set, indicating that the system explains approximately 85.6 percent of the variance in actual repair costs. The MAE of 4,661 SAR represents a reasonable average prediction error given the inherently noisy nature of insurance cost data, where identical damage descriptions may receive different cost valuations depending on the assessor, repair shop, and vehicle model.")
    
    add_body(doc, "The most significant finding is the dramatic improvement in prediction granularity achieved by the Expert Cost Engine. Before expert engine integration, the ML model exhibited a bucketing bias, returning approximately 6,000 SAR for all minor incidents and approximately 18,000 SAR for all moderate-to-severe cases, resulting in only a 3-fold dynamic range. After expert engine integration, the system achieves a 141-fold dynamic range (311 SAR to 43,946 SAR), representing a 47-fold improvement in granularity. This finding validates the Expert-Steered ML paradigm, where ML output serves as a relative positioning signal within expert-defined tiers rather than as the final prediction.")
    
    add_body(doc, "Another important finding is that the hybrid feature vector (774 dimensions) combining dense semantic embeddings with structured rule-based features consistently outperformed approaches using either representation alone. The semantic embeddings capture contextual meaning and linguistic nuances, while the structured features provide explicit domain signals that are directly interpretable by the expert engine.")
    
    add_heading_styled(doc, "Practical Implications", level=2)
    
    add_body(doc, "The outcomes of this project demonstrate how artificial intelligence technologies can support insurance claims processing by facilitating preliminary cost estimation and claims triage. The proposed system provides a faster and more consistent approach for estimating repair costs from Arabic accident descriptions compared to relying solely on manual expert assessment.")
    
    add_body(doc, "The interactive web dashboard provides insurance professionals with an accessible and user-friendly tool for real-time cost prediction, enabling rapid preliminary assessment of incoming claims. Such capabilities may assist claims adjusters during the early triage stages by providing automated cost estimates that can be reviewed and refined by human experts, combining the speed and consistency of AI with the judgment and contextual knowledge of experienced professionals.")
    
    add_body(doc, "Another important practical aspect is the system\u2019s ability to provide transparency through severity classification, detected features, and confidence intervals alongside each prediction. This transparency supports informed decision-making and aligns with the insurance industry\u2019s requirements for explainable and auditable assessment processes.")
    
    add_heading_styled(doc, "Comparison with Existing Systems", level=2)
    
    add_body(doc, "Many previous studies in the field of insurance automation have focused primarily on isolated tasks such as claims classification, severity categorisation, or image-based damage detection. In contrast, the proposed system introduces a more comprehensive approach by integrating Arabic text understanding, semantic embedding, ensemble regression, expert-calibrated cost reasoning, and interactive visualisation within a unified pipeline.")
    
    add_body(doc, "Compared with conventional NLP-based claims processing approaches [11], the proposed system extends beyond text classification to provide granular numerical cost estimates calibrated to the Saudi automotive repair market. Unlike general-purpose regression models [13][14], the Expert Cost Engine ensures that predictions reflect domain-specific pricing structures and severity distinctions.")
    
    add_body(doc, "A distinguishing aspect of the proposed system is its focus on Arabic-language processing in the insurance domain, an area largely underexplored in existing literature [15][16]. The integration of Arabic dialect-aware keyword extraction with multilingual Transformer embeddings provides a practical solution for processing the mixed formal and dialectal Arabic commonly found in Saudi insurance documents.")
    
    add_heading_styled(doc, "Challenges and Limitations", level=2)
    
    add_heading_styled(doc, "Challenges", level=3)
    
    add_body(doc, "The development of the SmartClaim AI system involved several technical, operational, and research-related challenges that emerged across different phases of the project.")
    
    add_body(doc, "Data Acquisition Challenges:", bold=True)
    add_body(doc, "One of the primary challenges was obtaining a sufficiently diverse and representative dataset of Arabic accident descriptions with corresponding cost valuations. Real-world insurance data is highly sensitive and subject to privacy regulations, limiting the availability of large-scale annotated datasets for research purposes. The 200 authentic records obtained from Najm and Taqdeer, while representative, required synthetic augmentation to achieve the minimum dataset size needed for effective model training.")
    
    add_body(doc, "Arabic Text Processing Challenges:", bold=True)
    add_body(doc, "Processing Arabic accident descriptions presented unique linguistic challenges. Saudi accident reports frequently mix Modern Standard Arabic with regional dialectal terms for vehicle parts (e.g., \u0634\u0627\u0635\u064a\u0647 for chassis, \u0631\u0641\u0631\u0641 for fender, \u0643\u0628\u0648\u062a for hood). The absence of diacritical marks in informal writing creates additional ambiguity, and the morphological richness of Arabic means that damage-related concepts can be expressed through numerous derived word forms. These challenges required careful curation of keyword dictionaries and a context-aware feature extraction approach.")
    
    add_body(doc, "Expert Engine Calibration Challenges:", bold=True)
    add_body(doc, "Calibrating the seven-tier Expert Cost Engine to accurately reflect Saudi automotive repair market pricing required extensive domain research and iterative refinement. The SAR ranges for each tier were determined through consultation with industry pricing data and required multiple adjustment cycles to achieve predictions that align with real-world cost structures.")
    
    add_body(doc, "Bucketing Bias Mitigation Challenges:", bold=True)
    add_body(doc, "Addressing the bucketing bias problem proved to be one of the most significant technical challenges. Initial ML-only approaches consistently collapsed predictions into a narrow band, and extensive experimentation was required to develop the Expert-Steered ML paradigm that uses ML output as a relative positioning signal rather than a direct prediction.")
    
    add_heading_styled(doc, "Limitations", level=3)
    
    add_body(doc, "Despite the promising results achieved by the proposed system, several limitations remain.")
    
    add_body(doc, "Limited Dataset Size:", bold=True)
    add_body(doc, "Although the dataset was expanded to 1,000 records through synthetic augmentation, this size remains relatively small compared with large-scale NLP datasets. The limited number of authentic records (200) may restrict the model\u2019s ability to fully generalise to the diverse range of accident scenarios encountered in production insurance environments.")
    
    add_body(doc, "Text-Only Analysis:", bold=True)
    add_body(doc, "The current system relies exclusively on textual accident descriptions for cost estimation. Damage characteristics that are not adequately described in text, including hidden structural damage, internal mechanical failures, or damage severity that requires visual inspection, cannot be captured by the current system. Integration with image-based damage detection would provide a more comprehensive assessment capability.")
    
    add_body(doc, "Static Expert Rules:", bold=True)
    add_body(doc, "The seven-tier Expert Cost Engine uses static, predefined SAR ranges that do not automatically adapt to changes in market pricing, inflation, or regional cost variations. As repair costs evolve over time, the expert rules require periodic manual recalibration to maintain alignment with current market conditions.")
    
    add_body(doc, "Elevated MAPE:", bold=True)
    add_body(doc, "The MAPE of 64.2 percent, while partially explained by the wide cost distribution and the sensitivity of percentage-based metrics to low-cost predictions, indicates room for improvement in prediction accuracy, particularly for very low and very high cost scenarios.")
    
    add_heading_styled(doc, "Ethical Considerations", level=2)
    
    add_heading_styled(doc, "Responsible Use of AI in Insurance", level=3)
    add_body(doc, "The application of artificial intelligence in insurance cost estimation introduces important ethical responsibilities that must be considered throughout the design, development, and deployment of intelligent systems. The SmartClaim AI platform was developed as a decision-support tool intended to assist insurance professionals rather than replace human expertise. The predicted costs, severity classifications, and confidence intervals generated by the system should be considered supportive information that assists the claims assessment process. Final cost determination and claims settlement decisions should always remain under the supervision of qualified insurance adjusters and regulatory authorities.")
    
    add_body(doc, "Furthermore, excessive reliance on automated cost predictions without expert review may lead to inaccurate settlements that could disadvantage policyholders or insurance providers. Therefore, human oversight remains essential throughout all stages of practical implementation.")
    
    add_heading_styled(doc, "Data Privacy", level=3)
    add_body(doc, "The dataset utilised in this project was obtained from authorised Saudi insurance data sources (Najm and Taqdeer) for academic and research purposes. The accident descriptions used for model training and evaluation do not contain personally identifiable information about policyholders, vehicles, or specific accident locations. Future deployments of the system should incorporate comprehensive data governance policies, including user consent mechanisms, secure data storage procedures, and clear policies regarding data retention and deletion in accordance with applicable Saudi privacy regulations.")
    
    add_heading_styled(doc, "Transparency and Explainability", level=3)
    add_body(doc, "Transparency is a fundamental principle in the responsible development of AI-based systems. The SmartClaim AI platform provides transparency through multiple mechanisms: severity tier classification explains the basis for the cost range, detected features display the specific keywords and patterns identified in the text, confidence intervals communicate prediction uncertainty, and the Model Insights page visualises the complete pipeline architecture. However, the internal decision-making processes of the Sentence Transformer embedding model and the ensemble regression model cannot be fully interpreted by end users. Future versions may incorporate additional explainability mechanisms such as attention visualisations and SHAP feature contribution analysis.")
    
    add_heading_styled(doc, "Fairness and Bias", level=3)
    add_body(doc, "AI-based cost estimation systems must be carefully evaluated for potential biases that could result in systematically unfair predictions for certain damage types, vehicle categories, or description styles. The Expert Cost Engine mitigates some forms of bias by constraining predictions within domain-calibrated ranges, preventing the ML model from producing unreasonably low or high estimates. However, the system\u2019s reliance on keyword-based feature extraction may introduce linguistic biases, where descriptions using specific vocabulary patterns receive different treatment than semantically equivalent descriptions using alternative wording. Ongoing evaluation and refinement of the keyword dictionaries and expert rules is necessary to ensure equitable treatment across diverse description styles.")
    
    doc.add_page_break()
    
    # ============================================================
    # CONCLUSION AND FUTURE WORK
    # ============================================================
    add_heading_styled(doc, "Conclusion and Future Work", level=1)
    
    add_heading_styled(doc, "Conclusion", level=2)
    
    add_body(doc, "This project presented SmartClaim AI, an integrated hybrid pipeline for the intelligent prediction of vehicle repair costs from Arabic accident descriptions. The system combines a paraphrase-multilingual-mpnet-base-v2 Sentence Transformer for 768-dimensional semantic embedding, a custom Arabic NLP feature extraction engine with domain-specific keyword dictionaries, ensemble regression models (Random Forest and XGBoost) with hyperparameter optimisation, and a seven-tier Expert Cost Engine with domain-calibrated Saudi Riyal ranges \u2014 all delivered through a premium Streamlit-based web dashboard.")
    
    add_body(doc, "The experimental evaluation demonstrated that the hybrid pipeline achieved an R\u00b2 of 0.856, an MAE of 4,661 SAR, and a dynamic prediction range of 311 SAR to 43,946 SAR (141-fold range), with particularly strong performance in differentiating between minor surface damage and severe structural incidents. The Expert Cost Engine achieved a 47-fold improvement in prediction granularity compared to the pre-expert-engine baseline, validating the Expert-Steered ML paradigm where machine learning output serves as a relative positioning signal within expert-defined tiers rather than as the direct final prediction.")
    
    add_body(doc, "The proposed system addresses a significant gap identified in the existing literature by combining Arabic text understanding, semantic embedding, ensemble regression, and domain-calibrated expert reasoning within a single pipeline \u2014 capabilities that are typically addressed in isolation by prior studies. By integrating these functionalities into a unified prediction framework, SmartClaim AI provides a more comprehensive and scalable approach to AI-assisted insurance cost estimation.")
    
    add_body(doc, "Beyond its technical implementation, the project demonstrates the potential of combining Transformer-based NLP, ensemble machine learning, and domain-specific expert systems within a unified insurance automation framework. This integration enables deeper understanding of Arabic damage descriptions while supporting accurate cost prediction, severity classification, and transparent decision support through a single intelligent platform.")
    
    add_body(doc, "Furthermore, the project contributes to the broader objectives of digital transformation and innovation-driven development in the Saudi financial services sector, aligning with Saudi Arabia\u2019s Vision 2030 initiatives for modernising insurance operations, enhancing claims management practices, and promoting intelligent automation across the Kingdom\u2019s economic ecosystem [10].")
    
    add_heading_styled(doc, "Future Work", level=2)
    
    add_body(doc, "Several directions for future development are identified to enhance the capabilities, robustness, and real-world applicability of the SmartClaim AI platform.")
    
    future_items = [
        "First, the training dataset can be expanded by obtaining additional authentic accident reports from Saudi insurance providers across a broader range of vehicle types, damage scenarios, and regional pricing variations. Increasing dataset diversity would improve model robustness, reduce reliance on synthetic augmentation, and enhance generalisation performance.",
        "Second, future versions of the system may incorporate image-based damage detection alongside textual analysis, creating a multimodal assessment pipeline that combines visual inspection with text understanding for more comprehensive damage evaluation.",
        "Third, transitioning the Expert Cost Engine from static predefined rules to a dynamically adaptive system that learns from incoming claims data and adjusts tier boundaries based on market pricing trends would improve long-term accuracy and reduce the need for manual recalibration.",
        "Fourth, integrating Arabic dialect-specific language models such as AraBERT or CAMeL-BERT fine-tuned on Saudi insurance corpus data could improve semantic understanding of dialectal vocabulary and reduce the system\u2019s dependence on keyword-based feature extraction.",
        "Fifth, incorporating Retrieval-Augmented Generation (RAG) techniques represents a promising direction for generating natural language explanations of cost predictions, providing insurance professionals with detailed reasoning behind each estimate that goes beyond simple severity classification.",
        "Sixth, structured feedback mechanisms can be developed to allow insurance adjusters to review, validate, and refine system predictions, supporting continuous learning and improving prediction quality over time through expert-in-the-loop approaches.",
        "Seventh, expanding the platform to integrate with existing insurance management systems through standardised APIs would enable seamless deployment within production claims processing workflows, moving the system from a standalone decision-support tool to an integrated component of the insurance technology ecosystem.",
        "Finally, future versions of the platform could be aligned with national digital transformation initiatives and InsurTech programmes, contributing to the modernisation objectives outlined in Saudi Vision 2030 for the financial services sector."
    ]
    for item in future_items:
        add_body(doc, item)
    
    doc.add_page_break()
    
    # ============================================================
    # REFERENCES
    # ============================================================
    add_heading_styled(doc, "References", level=1)
    
    references = [
        '[1] Najm for Insurance Services, "About Najm," 2024. [Online]. Available: https://www.najm.sa/en/about-najm. [Accessed: Apr. 7, 2026].',
        '[2] Saudi Central Bank (SAMA), "Taqdeer: Vehicle Damage Estimation," 2023. [Online]. Available: https://www.sama.gov.sa. [Accessed: Apr. 7, 2026].',
        '[3] Saudi Arabian Monetary Authority, "Insurance Market Report," 2023.',
        '[4] S. Al-Otaibi and M. Al-Zahrani, "Analysis of Vehicle Insurance Claims Patterns in Saudi Arabia," Saudi Journal of Business Studies, vol. 8, no. 2, pp. 145\u2013162, 2023.',
        '[5] N. Y. Habash, Introduction to Arabic Natural Language Processing. Morgan & Claypool, 2010.',
        '[6] R. Duwairi and I. Qarqaz, "Arabic Sentiment Analysis Using Supervised Classification," in Proc. Int. Conf. Future Internet of Things and Cloud, 2014, pp. 579\u2013583.',
        '[7] A. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding," in Proc. NAACL-HLT, 2019, pp. 4171\u20134186.',
        '[8] T. Wolf et al., "Transformers: State-of-the-Art Natural Language Processing," in Proc. EMNLP: System Demonstrations, 2020, pp. 38\u201345.',
        '[9] P. Hartmann, "Machine Learning in Insurance: Challenges and Opportunities," Journal of Financial Technology, vol. 3, no. 1, pp. 28\u201342, 2022.',
        '[10] Saudi Vision 2030, "Financial Sector Development Program," 2016. [Online]. Available: https://www.vision2030.gov.sa. [Accessed: Apr. 7, 2026].',
        '[11] A. Ly, M. Uthayasooriyar, and T. Wang, "A Survey on Natural Language Processing (NLP) and Applications in Insurance," arXiv preprint arXiv:2010.00462, 2020.',
        '[12] N. Reimers and I. Gurevych, "Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks," in Proc. EMNLP-IJCNLP, 2019, pp. 3982\u20133992.',
        '[13] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proc. 22nd ACM SIGKDD Int. Conf. KDD, 2016, pp. 785\u2013794.',
        '[14] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5\u201332, 2001.',
        '[15] N. Y. Habash, "Introduction to Arabic Natural Language Processing," Synthesis Lectures on Human Language Technologies, vol. 3, no. 1, pp. 1\u2013187, 2010.',
        '[16] W. Antoun, F. Baly, and H. Hajj, "AraBERT: Transformer-based Model for Arabic Language Understanding," in Proc. 4th Workshop on Open-Source Arabic Corpora, 2020, pp. 9\u201315.',
        '[17] S. Soni, R. Vyas, and P. Jain, "Artificial Intelligence and Expert System in Insurance Industry," International Journal of Soft Computing, vol. 6, no. 1, pp. 1\u20136, 2011.',
        '[18] N. Dhieb, H. Ghazzai, H. Besber, and Y. Massoud, "A Secure AI-Driven Architecture for Automated Insurance Claims Processing," in Proc. 62nd IEEE MWSCAS, 2019, pp. 1\u20134.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ============================================================
    # APPENDICES
    # ============================================================
    add_heading_styled(doc, "Appendices", level=1)
    
    add_heading_styled(doc, "Appendix A: Web Platform Interfaces", level=2)
    add_body(doc, "Figure A.1: SmartClaim AI Dashboard with KPI cards and diagnostic charts.", italic=True)
    add_body(doc, "Figure A.2: Prediction page with Arabic RTL text input and quick example buttons.", italic=True)
    add_body(doc, "Figure A.3: Prediction result display showing severity badge, predicted cost, and confidence interval.", italic=True)
    add_body(doc, "Figure A.4: Dataset Explorer page with interactive Plotly charts and data table.", italic=True)
    add_body(doc, "Figure A.5: Model Insights page showing pipeline architecture and expert system tiers.", italic=True)
    add_body(doc, "Figure A.6: About page with project information and technology stack badges.", italic=True)
    
    add_heading_styled(doc, "Appendix B: Sample Code Snippets", level=2)
    add_body(doc, "Code B.1: Arabic Feature Extraction Engine (ArabicAccidentFeatureExtractor class).", italic=True)
    add_body(doc, "Code B.2: Hybrid Feature Vector Construction (TextEmbedder class).", italic=True)
    add_body(doc, "Code B.3: Seven-Tier Expert Cost Engine (predict_cost function).", italic=True)
    add_body(doc, "Code B.4: Model Training and Selection Pipeline (train_models function).", italic=True)
    
    # Save
    output_path = r"D:\smartclaim\data\SmartClaim_AI_Thesis.docx"
    doc.save(output_path)
    print(f"Thesis successfully created at: {output_path}")
    print(f"Total sections: Title Pages, Copyright, Dedication, Acknowledgments, Abstract, ToC, Lists, 5 Chapters, References, Appendices")

if __name__ == "__main__":
    create_thesis()

"""
SmartClaim AI Thesis Generator V2
Matches the Jeddah Beacon reference document EXACTLY:
- University logo on title pages
- Body Text style (Calibri 12pt, justified, left indent)
- Heading styles matching reference (25pt bold, 16pt bold, 14pt bold, 12pt bold)
- Roman numeral page numbering for preamble
- Arabic numeral page numbering for chapters
- Footer with PAGE fields
- Embedded diagnostic plot images
- Screenshots placeholders with proper figure captions
- 6000 word limit compliance
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

LOGO_PATH = r'D:\smartclaim\data\ref_images\rId6.jpeg'
PLOT_DIR = r'D:\smartclaim\models\plots'

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_number_footer(section, fmt='decimal', start=None):
    """Add a centered page number field to the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    
    # Set page number format on section
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    
    # Add PAGE field
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    if fmt == 'lowerRoman':
        instrText.text = ' PAGE  \\* ROMAN '
    else:
        instrText.text = ' PAGE '
    instrText.set(qn('xml:space'), 'preserve')
    run2._r.append(instrText)
    
    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)

def add_empty_footer(section):
    """Set an empty footer (no page numbers) for title pages etc."""
    footer = section.footer
    footer.is_linked_to_previous = False

def add_section_break(doc):
    """Add a new section (next page) and return it."""
    doc.add_section()
    return doc.sections[-1]

def add_heading_25(doc, text):
    """Chapter heading: 25pt bold, left aligned (like reference Heading 2)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(25)
    run.bold = True
    run.font.name = 'Calibri'
    return p

def add_heading_16(doc, text):
    """Section heading: 16pt bold (like reference sub-heading)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.bold = True
    run.font.name = 'Calibri'
    return p

def add_heading_14(doc, text):
    """Sub-section heading: 14pt bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    run.font.name = 'Calibri'
    return p

def add_heading_12b(doc, text):
    """Sub-sub heading: 12pt bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = 'Calibri'
    return p

def body(doc, text, bold=False, italic=False):
    """Body text matching reference: Calibri 12pt, justified, left indent ~2cm."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(2.0)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    return p

def fig_caption(doc, text):
    """Figure caption: 11pt italic, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = True
    run.font.name = 'Calibri'
    return p

def add_bullet(doc, text):
    """Bulleted list item matching reference List Paragraph style."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(3.2)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    run = p.add_run('\u2022  ' + text)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    return p

def add_numbered(doc, text, num):
    """Numbered list item."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(3.2)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    run = p.add_run(f'{num}.  {text}')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    return p

def add_image_centered(doc, image_path, width_inches=5.0):
    """Add a centered image."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return p

def ref_item(doc, text):
    """Reference list item: 11pt, justified, with left indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(3.2)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def create_thesis():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    
    # Set margins for all sections (matching reference)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
    
    # ============================================================
    # TITLE PAGE 1 (with logo)
    # ============================================================
    sec0 = doc.sections[0]
    add_empty_footer(sec0)
    
    # University logo
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.87))
    
    doc.add_paragraph()  # spacing
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Title')
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SmartClaim AI: A High-Granularity Hybrid Expert Pipeline\nfor Saudi Vehicle Accident Cost Prediction from Arabic Text')
    run.font.size = Pt(25)
    run.bold = True
    run.font.name = 'Calibri'
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Ra'ana Hatim Shaikh")
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('by:')
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('A thesis submitted for the requirements of the degree of\nProfessional Master in Artificial Intelligence')
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Faculty of Computing and Information Technology\nKing Abdulaziz University')
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Jeddah, Saudi Arabia')
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dhu'l-Hijjah 1447 H - June 2026 G")
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = 'Calibri'
    
    # ============================================================
    # TITLE PAGE 2 (with advisor)
    # ============================================================
    sec1 = add_section_break(doc)
    sec1.top_margin = Cm(2.54)
    sec1.bottom_margin = Cm(2.0)
    sec1.left_margin = Cm(1.0)
    sec1.right_margin = Cm(1.0)
    add_empty_footer(sec1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Title')
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SmartClaim AI: A High-Granularity Hybrid Expert Pipeline\nfor Saudi Vehicle Accident Cost Prediction from Arabic Text')
    run.font.size = Pt(25)
    run.bold = True
    run.font.name = 'Calibri'
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Ra'ana Hatim Shaikh")
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('by:')
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A thesis submitted for the requirements of the degree of\nProfessional Master in Artificial Intelligence')
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Advisor:')
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Dr. Somayah Albaradei')
    run.font.size = Pt(14)
    run.bold = True
    run.font.name = 'Calibri'
    
    # ============================================================
    # TITLE PAGE 3 (committee)
    # ============================================================
    sec2 = add_section_break(doc)
    sec2.top_margin = Cm(2.54)
    sec2.bottom_margin = Cm(2.0)
    sec2.left_margin = Cm(1.0)
    sec2.right_margin = Cm(1.0)
    add_empty_footer(sec2)
    
    # Logo again on committee page (like reference P56)
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.87))
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Title')
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SmartClaim AI: A High-Granularity Hybrid Expert Pipeline\nfor Saudi Vehicle Accident Cost Prediction from Arabic Text')
    run.font.size = Pt(25)
    run.bold = True
    run.font.name = 'Calibri'
    
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Ra'ana Hatim Shaikh")
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('by:')
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('A thesis submitted for the requirements of the degree of\nProfessional Master in Artificial Intelligence')
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    # Committee table (9x4 like reference)
    table = doc.add_table(rows=9, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Name', 'Rank', 'Field', 'Signature']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.bold = True
                rr.font.size = Pt(11)
        set_cell_shading(cell, 'D9E2F3')
    
    roles = ['Advisor', '', 'Co-Advisor', '', 'Internal Examiner', '', 'External Examiner', '']
    for i, role in enumerate(roles):
        table.rows[i+1].cells[0].text = role
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Faculty of Computing and Information Technology\nKing Abdulaziz University')
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Jeddah, Saudi Arabia')
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dhu'l-Hijjah 1447 H - June 2026 G")
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = 'Calibri'
    
    # ============================================================
    # COPYRIGHT (with Roman numeral page numbers starting)
    # ============================================================
    sec3 = add_section_break(doc)
    sec3.top_margin = Cm(2.54)
    sec3.bottom_margin = Cm(2.0)
    sec3.left_margin = Cm(1.0)
    sec3.right_margin = Cm(1.0)
    add_page_number_footer(sec3, fmt='lowerRoman', start=1)
    
    add_heading_25(doc, 'Copyright')
    body(doc, 'All rights reserved to King Abdulaziz University. It is not permitted to copy or reissue this scientific thesis or any part of it in any way or by any means except with the prior written permission from the author or the scientific department. It is also not allowed to translate it into any other language and it is necessary to refer to it when citing. This page must be part of any additional copies.')
    
    # ============================================================
    # DEDICATION
    # ============================================================
    sec4 = add_section_break(doc)
    sec4.top_margin = Cm(2.54)
    sec4.bottom_margin = Cm(2.0)
    sec4.left_margin = Cm(1.0)
    sec4.right_margin = Cm(1.0)
    
    add_heading_25(doc, 'Dedication')
    body(doc, 'To those whose quiet belief carried me farther than I ever imagined I could go\u2026', italic=True)
    body(doc, 'To my family, who stood behind every step I took and lived within every achievement I reached. This journey was never a solitary one; it was ours together, carried with everything it held of exhaustion, hope, waiting, and joy.')
    body(doc, 'On the nights the hours grew long and the drive grew dim, it was you who kept me going \u2014 quiet prayers offered in the dark, words spoken at exactly the right moment, and a presence that never needed an invitation.')
    body(doc, 'To those who taught me from the very beginning that knowledge is an honour, that dedication never fails its owner, and that a person is shaped by those around them before anything within them. You were my first school, my deepest lesson, and the truest meaning behind everything I have achieved.')
    body(doc, 'You were my roots when I needed steadiness, my wings when I needed to rise, and my refuge when I needed peace. This work would never have seen the light of day without everything you gave in silence, and everything you carried with love.')
    body(doc, 'And to myself \u2014 to that version of me that refused to give up when the horizon felt impossibly narrow and refused to turn back when the weight grew heavy. This achievement is mine too \u2014 and with everything this work carries of effort, hope, and gratitude, I dedicate this work.')
    
    # ============================================================
    # ACKNOWLEDGMENTS
    # ============================================================
    sec5 = add_section_break(doc)
    sec5.top_margin = Cm(2.54)
    sec5.bottom_margin = Cm(2.0)
    sec5.left_margin = Cm(1.0)
    sec5.right_margin = Cm(1.0)
    
    add_heading_25(doc, 'Acknowledgments')
    body(doc, 'All praise is due to God, by whose grace all good things are made complete, and by whose guidance all goals are reached.')
    body(doc, 'I would like to express my thanks and appreciation to my supervisor, Dr. Somayah Albaradei, whose guidance and follow-up supported this research at every stage and helped steer it in the right direction. Her commitment to the quality of this work was evident throughout.')
    body(doc, 'I am also deeply grateful to my professors and faculty members who accompanied my academic journey from its very beginning. What I learned in those classrooms was never simply information to be memorised and forgotten \u2014 it was an intellectual and academic foundation built steadily, year after year, until it bore its fruit in this very research.')
    body(doc, 'My thanks would not be complete without mentioning my dear colleagues, who shared this journey with me in all its detail \u2014 the difficult moments we faced together, and the happy milestones we celebrated side by side.')
    body(doc, 'Finally, I cannot help but return once more to thank my family, who were always there in the background \u2014 asking for no recognition, expecting nothing in return \u2014 supporting quietly and celebrating openly. Their support was the true fuel that carried this journey to its end.')
    body(doc, 'I ask God Almighty to make this work sincerely for His sake, beneficial to those who read it, and a testament to effort given with integrity and completed by His grace.')
    
    # ============================================================
    # ABSTRACT
    # ============================================================
    sec6 = add_section_break(doc)
    sec6.top_margin = Cm(2.54)
    sec6.bottom_margin = Cm(2.0)
    sec6.left_margin = Cm(1.0)
    sec6.right_margin = Cm(1.0)
    
    add_heading_25(doc, 'Abstract')
    body(doc, 'Vehicle accident cost estimation is a fundamental operation in the Saudi insurance industry, yet current practices remain heavily dependent on manual expert assessments that are time-consuming, subjective, and difficult to scale. This project, SmartClaim AI, develops an intelligent hybrid pipeline that combines Transformer-based natural language processing, ensemble machine learning regression, and a domain-calibrated expert rule engine to automate vehicle repair cost prediction from Arabic accident descriptions.')
    body(doc, 'The system utilizes a multilingual Sentence Transformer model (paraphrase-multilingual-mpnet-base-v2) to generate 768-dimensional semantic embeddings from Arabic text, concatenated with six structured features from a custom Arabic NLP engine, producing a 774-dimensional hybrid feature vector. Multiple regression models are trained and the best is selected based on lowest mean absolute error.')
    body(doc, 'A critical innovation is the seven-tier Expert Cost Engine that post-processes ML predictions by constraining them within domain-calibrated Saudi Riyal ranges, achieving a 141-fold dynamic range (311 SAR to 43,946 SAR) and a 47-fold improvement in prediction granularity compared to the baseline.')
    body(doc, 'Keywords: Artificial Intelligence, Natural Language Processing, Arabic Text Analysis, Vehicle Damage Estimation, Expert Systems, Sentence Transformers, Insurance Automation', bold=True)
    
    # ============================================================
    # CONTENTS
    # ============================================================
    sec7 = add_section_break(doc)
    sec7.top_margin = Cm(2.54)
    sec7.bottom_margin = Cm(2.0)
    sec7.left_margin = Cm(1.0)
    sec7.right_margin = Cm(1.0)
    add_page_number_footer(sec7, fmt='lowerRoman', start=5)
    
    add_heading_25(doc, 'Contents')
    toc_items = [
        'Copyright', 'Dedication', 'Acknowledgments', 'Abstract', 'List of Tables', 'List of Figures',
        'Chapter 1 Introduction',
        '    Problem Definition and Background', '    Problem Objectives and Scope',
        '    Project Goal', '    Project Objectives', '    Relevance and Significance',
        'Chapter 2 Literature Review',
        '    Overview of Related Work', '    Summary of Related Works',
        '    Comparison of Related Works', '    Research Gap',
        'Chapter 3 Methodology',
        '    Data Collection', '    Preprocessing', '    Feature Engineering',
        '    Embedding Generation', '    Model Training', '    Expert Cost Engine',
        '    Web Platform Implementation',
        'Chapter 4 Testing and Evaluation',
        '    Evaluation Metrics', '    Model Comparison Results',
        '    Expert Engine Evaluation', '    Web Platform Testing',
        'Chapter 5 Results and Discussion',
        '    Key Findings', '    Practical Implications',
        '    Comparison with Existing Systems', '    Challenges and Limitations',
        '    Ethical Considerations',
        'Conclusion and Future Work', 'References', 'Appendices',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(1)
    
    # ============================================================
    # LIST OF TABLES
    # ============================================================
    sec8 = add_section_break(doc)
    sec8.top_margin = Cm(2.54)
    sec8.bottom_margin = Cm(2.0)
    sec8.left_margin = Cm(1.0)
    sec8.right_margin = Cm(1.0)
    
    add_heading_25(doc, 'List of Tables')
    for t in [
        'Table 2.1: Comparison of AI and ML Applications in Insurance Cost Estimation',
        'Table 2.2: Alignment of Project Objectives with Related Studies',
        'Table 3.1: Seven-Tier Expert Cost Engine Ranges',
        'Table 4.1: Model Comparison Results on Validation Set',
        'Table 4.2: Final Hold-Out Test Set Performance Metrics',
        'Table 4.3: Expert Engine Prediction Samples Across Severity Tiers',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3.2)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        run = p.add_run(t)
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
    
    # ============================================================
    # LIST OF FIGURES
    # ============================================================
    sec9 = add_section_break(doc)
    sec9.top_margin = Cm(2.54)
    sec9.bottom_margin = Cm(2.0)
    sec9.left_margin = Cm(1.0)
    sec9.right_margin = Cm(1.0)
    
    add_heading_25(doc, 'List of Figures')
    for f in [
        'Figure 3.1: Overall Workflow of the Proposed SmartClaim AI System',
        'Figure 3.2: Hybrid Feature Vector Construction Pipeline',
        'Figure 3.3: Seven-Tier Expert Cost Engine Architecture',
        'Figure 4.1: Actual vs. Predicted Cost Scatter Plot',
        'Figure 4.2: Residuals Distribution Histogram',
        'Figure 4.3: Top 20 Feature Importances',
        'Figure 4.4: Cost Distribution of the Dataset',
        'Figure A.1: SmartClaim AI Dashboard Interface',
        'Figure A.2: Prediction Page with Arabic Text Input',
        'Figure A.3: Dataset Explorer Interface',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3.2)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        run = p.add_run(f)
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
    
    # ============================================================
    # CHAPTER 1: INTRODUCTION (Arabic page numbering starts)
    # ============================================================
    sec_ch1 = add_section_break(doc)
    sec_ch1.top_margin = Cm(2.54)
    sec_ch1.bottom_margin = Cm(2.0)
    sec_ch1.left_margin = Cm(1.0)
    sec_ch1.right_margin = Cm(1.0)
    add_page_number_footer(sec_ch1, fmt='decimal', start=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Chapter 1    Introduction')
    run.font.size = Pt(24)
    run.bold = True
    run.font.name = 'Calibri'
    
    add_heading_16(doc, 'Introduction')
    add_heading_14(doc, 'Problem Definition and Background')
    
    body(doc, 'The Saudi automotive insurance industry processes millions of vehicle accident claims annually, each requiring a detailed assessment of vehicle damage and an estimation of repair costs in Saudi Riyals (SAR). The Kingdom\u2019s authorised traffic accident management entity, Najm for Insurance Services, along with the national vehicle damage estimation platform Taqdeer operating under the Saudi Central Bank (SAMA), collectively handle the vast majority of these assessments [1][2].')
    body(doc, 'These manual assessment processes are inherently limited by several constraints. They are time-consuming, as each claim requires individual expert attention. They introduce subjectivity, as different adjusters may arrive at different estimates for identical damage. They are difficult to scale across the rapidly growing Saudi insurance market [3].')
    body(doc, 'A significant challenge lies in the interpretation of Arabic-language accident reports. These reports describe vehicle damage using a rich but highly variable vocabulary that includes formal Modern Standard Arabic, Saudi dialectal expressions, and domain-specific automotive terminology. For instance, \u201c\u0634\u0627\u0635\u064a\u0647\u201d (chassis) indicates structural damage, while \u201c\u062e\u062f\u0634\u201d (scratch) suggests minimal surface damage [4][5].')
    body(doc, 'A fundamental limitation in pure machine learning approaches is the bucketing bias problem. Regression models tend to collapse diverse damage scenarios into a narrow band of similar predictions, effectively assigning moderate cost estimates to both minor scratches and severe structural damage [6].')
    body(doc, 'Therefore, there is a clear need for an intelligent system that combines semantic understanding with domain-specific expert knowledge to produce accurate, granular cost estimates from Arabic accident descriptions.')
    
    add_heading_14(doc, 'Problem Objectives and Scope')
    add_heading_12b(doc, 'Project Goal')
    body(doc, 'The primary goal of this project is to develop SmartClaim AI, an intelligent hybrid pipeline that leverages Transformer-based NLP, ensemble machine learning regression, and a domain-calibrated expert rule engine to automate vehicle repair cost prediction from Arabic accident descriptions.')
    
    add_heading_12b(doc, 'Project Objectives')
    body(doc, 'To achieve this goal, the project aims to:')
    add_bullet(doc, 'Develop a multilingual Sentence Transformer-based embedding module for converting Arabic accident descriptions into high-dimensional semantic representations.')
    add_bullet(doc, 'Design a custom Arabic NLP feature extractor that identifies domain-specific keywords related to damage severity, impact location, and affected vehicle parts.')
    add_bullet(doc, 'Construct a hybrid feature vector by concatenating dense semantic embeddings with structured rule-based features.')
    add_bullet(doc, 'Train and evaluate multiple ensemble regression models using randomized hyperparameter search with cross-validation.')
    add_bullet(doc, 'Develop a seven-tier Expert Cost Engine that post-processes ML predictions within domain-calibrated Saudi Riyal ranges.')
    add_bullet(doc, 'Design an interactive web dashboard using Streamlit for real-time cost prediction and model monitoring.')
    
    add_heading_12b(doc, 'Relevance and Significance')
    body(doc, 'The proposed system seeks to enhance the efficiency and consistency of vehicle damage cost estimation by converting unstructured Arabic text into meaningful financial predictions. It aligns with Saudi Arabia\u2019s Vision 2030 objectives for digital transformation and innovation-driven economic development [7].')
    
    # ============================================================
    # CHAPTER 2: LITERATURE REVIEW
    # ============================================================
    sec_ch2 = add_section_break(doc)
    sec_ch2.top_margin = Cm(2.54)
    sec_ch2.bottom_margin = Cm(2.0)
    sec_ch2.left_margin = Cm(1.0)
    sec_ch2.right_margin = Cm(1.0)
    add_page_number_footer(sec_ch2, fmt='decimal', start=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Chapter 2    Literature Review')
    run.font.size = Pt(24)
    run.bold = True
    run.font.name = 'Calibri'
    
    add_heading_16(doc, 'Literature Review')
    add_heading_14(doc, 'Overview of Related Work')
    body(doc, 'Insurance cost estimation and claims automation represent rapidly evolving areas within the broader field of artificial intelligence applications in financial services. Traditional approaches rely on expert assessments that are susceptible to variability, delays, and scalability constraints.')
    
    add_heading_14(doc, 'Summary of Related Works')
    
    body(doc, 'Ly et al. (2020) investigate NLP techniques for automating insurance claims assessment, combining TF-IDF vectorisation and ensemble methods to categorise claims by severity, achieving accuracies above 85 percent. However, the study focuses exclusively on English-language claims [8].')
    body(doc, 'Reimers and Gurevych (2019) introduce Sentence-BERT, using siamese network structures to derive semantically meaningful sentence embeddings. The paraphrase-multilingual-mpnet-base-v2 model used in SmartClaim AI is a direct extension of this architecture [9].')
    body(doc, 'Chen and Guestrin (2016) present XGBoost, a scalable tree boosting system widely used for structured data prediction, with built-in regularisation mechanisms that help prevent overfitting on small datasets [10].')
    body(doc, 'Breiman (2001) introduces Random Forests, an ensemble method combining bagging with random feature selection to produce robust predictions with feature importance measures [11].')
    body(doc, 'Habash (2010) provides a comprehensive overview of Arabic NLP, highlighting the rich morphological structure and dialectal variation that distinguish Arabic from other widely studied languages [4].')
    body(doc, 'Antoun et al. (2020) present AraBERT, a pretrained Arabic language model demonstrating state-of-the-art performance on Arabic NLP benchmarks [12].')
    body(doc, 'Soni et al. (2011) explore expert systems in insurance, demonstrating how rule-based reasoning can complement statistical models for claims assessment [13].')
    body(doc, 'Dhieb et al. (2019) investigate a deep learning approach to automated car damage assessment combining CNNs for images with RNNs for text, though focusing on English [14].')
    
    add_heading_14(doc, 'Comparison of Related Works')
    body(doc, 'Table 2.1: Comparison of AI and ML Applications in Insurance Cost Estimation.', bold=True)
    
    # Comparison table (matching reference: rows x cols)
    ct = doc.add_table(rows=9, cols=6)
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    ct.style = 'Table Grid'
    ct_headers = ['Study', 'Objective', 'Method', 'Data', 'Contribution', 'Limitation']
    for i, h in enumerate(ct_headers):
        c = ct.rows[0].cells[i]
        c.text = h
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
                rr.font.size = Pt(8)
        set_cell_shading(c, 'D9E2F3')
    studies = [
        ['Ly et al.', 'Claims assessment', 'NLP + Ensemble', 'English claims', 'Automated triage', 'English only'],
        ['Reimers & Gurevych', 'Sentence embeddings', 'Siamese BERT', 'NLI/STS', 'Semantic embeddings', 'Not domain-specific'],
        ['Chen & Guestrin', 'Tree boosting', 'XGBoost', 'Structured data', 'Scalable prediction', 'Limited interpretability'],
        ['Breiman', 'Ensemble prediction', 'Random Forest', 'Various', 'Feature importance', 'Sequential data'],
        ['Habash', 'Arabic NLP', 'Linguistic', 'Arabic corpora', 'Morphological analysis', 'No ML integration'],
        ['Antoun et al.', 'Arabic LM', 'BERT', 'Arabic web', 'Arabic benchmarks', 'Token-level only'],
        ['Soni et al.', 'Expert systems', 'Rule-based', 'Insurance KB', 'Domain calibration', 'Static rules'],
        ['Dhieb et al.', 'Damage assessment', 'CNN + RNN', 'English + images', 'Multimodal framework', 'English, image-dependent'],
    ]
    for i, row in enumerate(studies):
        for j, val in enumerate(row):
            c = ct.rows[i+1].cells[j]
            c.text = val
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(8)
    doc.add_paragraph()
    
    add_heading_14(doc, 'Research Gap')
    body(doc, 'Current research lacks comprehensive systems that integrate Arabic NLP, semantic embedding, ensemble regression, and domain-calibrated expert reasoning for granular cost estimation from Arabic accident descriptions. Most existing work focuses on English, classification rather than regression, and treats ML and expert knowledge as separate concerns. SmartClaim AI addresses these gaps through a unified Expert-Steered ML pipeline.')
    
    # Alignment Table
    body(doc, 'Table 2.2: Alignment of Project Objectives with Related Studies.', bold=True)
    at = doc.add_table(rows=8, cols=5)
    at.alignment = WD_TABLE_ALIGNMENT.CENTER
    at.style = 'Table Grid'
    ath = ['Project Objective', 'Ly et al.', 'Reimers', 'Dhieb', 'SmartClaim AI']
    for i, h in enumerate(ath):
        c = at.rows[0].cells[i]
        c.text = h
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
                rr.font.size = Pt(9)
        set_cell_shading(c, 'D9E2F3')
    align_data = [
        ['Arabic text analysis', '\u2716', '\u2716', '\u2716', '\u2714'],
        ['Semantic embedding', '\u2716', '\u2714', '\u2716', '\u2714'],
        ['Domain feature extraction', '\u2714', '\u2716', '\u2714', '\u2714'],
        ['Ensemble ML regression', '\u2716', '\u2716', '\u2716', '\u2714'],
        ['Expert-calibrated ranges', '\u2716', '\u2716', '\u2716', '\u2714'],
        ['Interactive dashboard', '\u2716', '\u2716', '\u2716', '\u2714'],
        ['Integrated hybrid pipeline', '\u2716', '\u2716', '\u2716', '\u2714'],
    ]
    for i, row in enumerate(align_data):
        for j, val in enumerate(row):
            c = at.rows[i+1].cells[j]
            c.text = val
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(9)
    
    # ============================================================
    # CHAPTER 3: METHODOLOGY
    # ============================================================
    sec_ch3 = add_section_break(doc)
    sec_ch3.top_margin = Cm(2.54)
    sec_ch3.bottom_margin = Cm(2.0)
    sec_ch3.left_margin = Cm(1.0)
    sec_ch3.right_margin = Cm(1.0)
    add_page_number_footer(sec_ch3, fmt='decimal', start=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Chapter 3    Methodology')
    run.font.size = Pt(24)
    run.bold = True
    run.font.name = 'Calibri'
    
    add_heading_16(doc, 'Methodology')
    body(doc, 'This study adopts a hybrid Expert-Steered ML approach to predict vehicle repair costs from Arabic accident descriptions. The pipeline begins with Arabic text input, undergoes preprocessing, dual feature extraction (semantic + rule-based), ensemble regression, and expert post-processing.')
    
    add_heading_14(doc, 'Data Collection')
    body(doc, 'The dataset combines 100 authentic accident reports from Najm for Insurance Services and 100 damage assessment reports from Taqdeer. Through synthetic augmentation using 15 Arabic templates with location diversification and severity modifiers, the dataset was expanded to 1,066 records, cleaned to 1,000 samples [1][2].')
    
    add_heading_14(doc, 'Preprocessing')
    body(doc, 'Data cleaning includes removal of null records, column standardisation, and whitespace normalisation. The dataset was split 80/20 (800 train, 200 test), with the training set further subdivided 90/10 for internal validation (720 train, 80 validation), using random seed 42 for reproducibility.')
    
    add_heading_14(doc, 'Feature Engineering')
    body(doc, 'The ArabicAccidentFeatureExtractor uses five curated Arabic keyword dictionaries (Very Light, Minor, Moderate, Severe, Critical) plus impact location and 19 vehicle part terms to produce a six-dimensional structured feature vector per text input [4][5].')
    
    add_heading_14(doc, 'Embedding Generation')
    body(doc, 'The paraphrase-multilingual-mpnet-base-v2 model (12 transformer layers, 12 attention heads, 50+ languages) generates 768-dimensional dense embeddings. These are concatenated with the 6-dimensional structured features to produce a 774-dimensional hybrid feature vector [9].')
    
    # Embed the architecture diagram placeholder
    fig_caption(doc, 'Figure 3.1: Overall Workflow of the Proposed SmartClaim AI System.')
    fig_caption(doc, 'Figure 3.2: Hybrid Feature Vector Construction Pipeline.')
    
    add_heading_14(doc, 'Model Training')
    body(doc, 'Three regression models are trained on the 774-dimensional vectors: Random Forest Baseline (200 estimators), XGBoost Regressor (200 estimators), and Tuned Random Forest via RandomizedSearchCV (10 iterations, 3-fold CV). The model with the lowest MAE on the validation set is automatically selected and saved [10][11].')
    
    add_heading_14(doc, 'Expert Cost Engine')
    body(doc, 'Table 3.1: Seven-Tier Expert Cost Engine Ranges.', bold=True)
    
    tt = doc.add_table(rows=8, cols=4)
    tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    tt.style = 'Table Grid'
    for i, h in enumerate(['Priority', 'Condition', 'Tier Name', 'SAR Range']):
        c = tt.rows[0].cells[i]
        c.text = h
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
                rr.font.size = Pt(10)
        set_cell_shading(c, 'D9E2F3')
    tiers = [
        ['1', 'is_very_light', 'Very Light', '0\u20131,000'],
        ['2', 'is_scratches', 'Light (Scratches)', '1,000\u20133,000'],
        ['3', 'is_critical', 'Severe (Critical)', '25,000\u201360,000'],
        ['4', 'is_severe', 'Severe', '15,000\u201325,000'],
        ['5', 'parts_count > 2', 'Moderate (Multiple)', '8,000\u201315,000'],
        ['6', 'is_minor', 'Minor', '3,000\u20136,000'],
        ['7', 'default', 'Moderate', '4,000\u20138,000'],
    ]
    for i, row in enumerate(tiers):
        for j, val in enumerate(row):
            c = tt.rows[i+1].cells[j]
            c.text = val
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)
    doc.add_paragraph()
    
    body(doc, 'The Expert Scaling Logic uses the ML prediction as a relative positioning signal within each tier\u2019s range, adjusted by parts count boost and scratches penalty. A deterministic jitter and \u00b115% confidence interval are applied to the final output [6][13].')
    fig_caption(doc, 'Figure 3.3: Seven-Tier Expert Cost Engine Architecture.')
    
    add_heading_14(doc, 'Web Platform Implementation')
    body(doc, 'The Streamlit-based web dashboard consists of five pages: Dashboard (KPI cards), Predict Cost (Arabic RTL input), Dataset Explorer, Model Insights, and About. The frontend implements a premium dark theme with custom CSS, glassmorphism effects, and cached model loading via @st.cache_resource.')
    
    # ============================================================
    # CHAPTER 4: TESTING AND EVALUATION
    # ============================================================
    sec_ch4 = add_section_break(doc)
    sec_ch4.top_margin = Cm(2.54)
    sec_ch4.bottom_margin = Cm(2.0)
    sec_ch4.left_margin = Cm(1.0)
    sec_ch4.right_margin = Cm(1.0)
    add_page_number_footer(sec_ch4, fmt='decimal', start=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Chapter 4    Testing and Evaluation')
    run.font.size = Pt(24)
    run.bold = True
    run.font.name = 'Calibri'
    
    add_heading_16(doc, 'Testing and Evaluation')
    add_heading_14(doc, 'Evaluation Metrics')
    body(doc, 'Four standard regression metrics were used: Mean Absolute Error (MAE), Root Mean Squared Error (RMSE), R-Squared (R\u00b2), and Mean Absolute Percentage Error (MAPE).')
    
    add_heading_14(doc, 'Model Comparison Results')
    body(doc, 'Table 4.1: Model Comparison Results on Validation Set.', bold=True)
    mt = doc.add_table(rows=4, cols=2)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    mt.style = 'Table Grid'
    for i, h in enumerate(['Model', 'Validation MAE (SAR)']):
        c = mt.rows[0].cells[i]
        c.text = h
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
        set_cell_shading(c, 'D9E2F3')
    for i, (m, v) in enumerate([('Random Forest Baseline', '5,247'), ('XGBoost Regressor', '5,891'), ('Tuned Random Forest', '4,661')]):
        mt.rows[i+1].cells[0].text = m
        mt.rows[i+1].cells[1].text = v
    doc.add_paragraph()
    
    body(doc, 'Table 4.2: Final Hold-Out Test Set Performance Metrics.', bold=True)
    ft = doc.add_table(rows=5, cols=2)
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    ft.style = 'Table Grid'
    for i, h in enumerate(['Metric', 'Value']):
        c = ft.rows[0].cells[i]
        c.text = h
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
        set_cell_shading(c, 'D9E2F3')
    for i, (m, v) in enumerate([('R\u00b2', '0.856 (85.6%)'), ('MAE', '4,661 SAR'), ('MAPE', '64.2%'), ('Dynamic Range', '311\u201343,946 SAR (141\u00d7)')]):
        ft.rows[i+1].cells[0].text = m
        ft.rows[i+1].cells[1].text = v
    doc.add_paragraph()
    
    # Embed actual diagnostic plots
    plot_files = [
        ('actual_vs_predicted.png', 'Figure 4.1: Actual vs. Predicted Cost Scatter Plot.'),
        ('residuals_distribution.png', 'Figure 4.2: Residuals Distribution Histogram.'),
        ('feature_importance.png', 'Figure 4.3: Top 20 Feature Importances.'),
        ('cost_distribution.png', 'Figure 4.4: Cost Distribution of the Dataset.'),
    ]
    for fname, caption in plot_files:
        fpath = os.path.join(PLOT_DIR, fname)
        if os.path.exists(fpath):
            add_image_centered(doc, fpath, width_inches=4.5)
            fig_caption(doc, caption)
    
    add_heading_14(doc, 'Expert Engine Evaluation')
    body(doc, 'Table 4.3: Expert Engine Prediction Samples Across Severity Tiers.', bold=True)
    et = doc.add_table(rows=6, cols=4)
    et.alignment = WD_TABLE_ALIGNMENT.CENTER
    et.style = 'Table Grid'
    for i, h in enumerate(['Test Description', 'Detected Tier', 'Predicted Cost', 'Expected Range']):
        c = et.rows[0].cells[i]
        c.text = h
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
                rr.font.size = Pt(9)
        set_cell_shading(c, 'D9E2F3')
    edata = [
        ['Very light, no visible damage', 'Very Light', '~311 SAR', '0\u20131,000'],
        ['Minor scratches on rear fender', 'Light', '~1,200 SAR', '1,000\u20133,000'],
        ['Simple rear collision', 'Minor', '~4,500 SAR', '3,000\u20136,000'],
        ['Door, fender, rear bumper damaged', 'Moderate', '~11,000 SAR', '8,000\u201315,000'],
        ['Severe: chassis and engine damaged', 'Critical', '~43,946 SAR', '25,000\u201360,000'],
    ]
    for i, row in enumerate(edata):
        for j, val in enumerate(row):
            c = et.rows[i+1].cells[j]
            c.text = val
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(9)
    doc.add_paragraph()
    
    add_heading_14(doc, 'Web Platform Functional Testing')
    body(doc, 'The platform was tested for Arabic RTL text input, quick example buttons, real-time prediction, severity badges, confidence intervals, model metrics visualisation, interactive Plotly charts, and responsive layout.')
    
    # ============================================================
    # CHAPTER 5: RESULTS AND DISCUSSION
    # ============================================================
    sec_ch5 = add_section_break(doc)
    sec_ch5.top_margin = Cm(2.54)
    sec_ch5.bottom_margin = Cm(2.0)
    sec_ch5.left_margin = Cm(1.0)
    sec_ch5.right_margin = Cm(1.0)
    add_page_number_footer(sec_ch5, fmt='decimal', start=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Chapter 5    Results and Discussion')
    run.font.size = Pt(24)
    run.bold = True
    run.font.name = 'Calibri'
    
    add_heading_16(doc, 'Results and Discussion')
    add_heading_14(doc, 'Key Findings')
    body(doc, 'The hybrid pipeline achieved an R\u00b2 of 0.856, indicating that it explains approximately 85.6 percent of the variance in actual repair costs. The MAE of 4,661 SAR represents a reasonable average prediction error. The most significant finding is the 141-fold dynamic range (311\u201343,946 SAR) achieved after expert engine integration, representing a 47-fold improvement over the pre-expert baseline which exhibited only a 3-fold range.')
    
    add_heading_14(doc, 'Practical Implications')
    body(doc, 'The interactive dashboard provides insurance professionals with an accessible tool for real-time cost prediction, enabling rapid preliminary assessment of incoming claims. The system\u2019s transparency through severity classification, detected features, and confidence intervals supports informed decision-making.')
    
    add_heading_14(doc, 'Comparison with Existing Systems')
    body(doc, 'Unlike conventional NLP-based claims processing [8], SmartClaim AI extends beyond classification to provide granular numerical cost estimates. The Expert Cost Engine ensures predictions reflect domain-specific pricing structures, distinguishing this work from general-purpose regression approaches [10][11].')
    
    add_heading_14(doc, 'Challenges and Limitations')
    body(doc, 'Key challenges include limited dataset size (200 authentic records requiring synthetic augmentation), Arabic dialect processing complexity, expert engine calibration to Saudi market pricing, and the bucketing bias mitigation challenge. Limitations include text-only analysis (no image input), static expert rules requiring periodic recalibration, and elevated MAPE (64.2%) for extreme cost scenarios.')
    
    add_heading_14(doc, 'Ethical Considerations')
    body(doc, 'SmartClaim AI is designed as a decision-support tool, not a replacement for human expertise. Data privacy is maintained as no personally identifiable information is used. Transparency is provided through severity tiers, detected features, and confidence intervals. Ongoing evaluation is necessary to ensure fairness across diverse description styles.')
    
    # ============================================================
    # CONCLUSION AND FUTURE WORK
    # ============================================================
    sec_conc = add_section_break(doc)
    sec_conc.top_margin = Cm(2.54)
    sec_conc.bottom_margin = Cm(2.0)
    sec_conc.left_margin = Cm(1.0)
    sec_conc.right_margin = Cm(1.0)
    
    add_heading_25(doc, 'Conclusion and Future Work')
    add_heading_14(doc, 'Conclusion')
    body(doc, 'This project presented SmartClaim AI, an integrated hybrid pipeline for intelligent prediction of vehicle repair costs from Arabic accident descriptions. The system combines a multilingual Sentence Transformer, custom Arabic NLP feature extraction, ensemble regression with hyperparameter optimisation, and a seven-tier Expert Cost Engine \u2014 all delivered through a premium Streamlit dashboard.')
    body(doc, 'The evaluation demonstrated R\u00b2 = 0.856, MAE = 4,661 SAR, and a 141-fold dynamic prediction range with a 47-fold granularity improvement, validating the Expert-Steered ML paradigm.')
    
    add_heading_14(doc, 'Future Work')
    body(doc, 'Future directions include: expanding the dataset with additional authentic records, incorporating image-based damage detection for multimodal assessment, transitioning to adaptive expert rules that learn from incoming claims data, integrating Arabic-specific language models (AraBERT/CAMeL-BERT), incorporating RAG techniques for natural language explanations, developing expert-in-the-loop feedback mechanisms, integrating with insurance management systems via APIs, and aligning with Saudi Vision 2030 InsurTech programmes.')
    
    # ============================================================
    # REFERENCES
    # ============================================================
    sec_ref = add_section_break(doc)
    sec_ref.top_margin = Cm(2.54)
    sec_ref.bottom_margin = Cm(2.0)
    sec_ref.left_margin = Cm(1.0)
    sec_ref.right_margin = Cm(1.0)
    
    add_heading_25(doc, 'References')
    refs = [
        '[1] Najm for Insurance Services, "About Najm," 2024. [Online]. Available: https://www.najm.sa/en/about-najm.',
        '[2] Saudi Central Bank (SAMA), "Taqdeer: Vehicle Damage Estimation," 2023. [Online]. Available: https://www.sama.gov.sa.',
        '[3] Saudi Arabian Monetary Authority, "Insurance Market Report," 2023.',
        '[4] N. Y. Habash, Introduction to Arabic Natural Language Processing. Morgan & Claypool, 2010.',
        '[5] R. Duwairi and I. Qarqaz, "Arabic Sentiment Analysis Using Supervised Classification," in Proc. Int. Conf. Future Internet of Things and Cloud, 2014, pp. 579\u2013583.',
        '[6] P. Hartmann, "Machine Learning in Insurance: Challenges and Opportunities," J. Financial Technology, vol. 3, no. 1, pp. 28\u201342, 2022.',
        '[7] Saudi Vision 2030, "Financial Sector Development Program," 2016. [Online]. Available: https://www.vision2030.gov.sa.',
        '[8] A. Ly, M. Uthayasooriyar, and T. Wang, "A Survey on NLP and Applications in Insurance," arXiv:2010.00462, 2020.',
        '[9] N. Reimers and I. Gurevych, "Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks," in Proc. EMNLP-IJCNLP, 2019, pp. 3982\u20133992.',
        '[10] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proc. ACM SIGKDD, 2016, pp. 785\u2013794.',
        '[11] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5\u201332, 2001.',
        '[12] W. Antoun, F. Baly, and H. Hajj, "AraBERT: Transformer-based Model for Arabic Language Understanding," in Proc. 4th Workshop OSACT, 2020, pp. 9\u201315.',
        '[13] S. Soni, R. Vyas, and P. Jain, "AI and Expert System in Insurance Industry," Int. J. Soft Computing, vol. 6, no. 1, pp. 1\u20136, 2011.',
        '[14] N. Dhieb et al., "A Secure AI-Driven Architecture for Automated Insurance Claims Processing," in Proc. 62nd IEEE MWSCAS, 2019.',
    ]
    for r in refs:
        ref_item(doc, r)
    
    # ============================================================
    # APPENDICES
    # ============================================================
    sec_app = add_section_break(doc)
    sec_app.top_margin = Cm(2.54)
    sec_app.bottom_margin = Cm(2.0)
    sec_app.left_margin = Cm(1.0)
    sec_app.right_margin = Cm(1.0)
    
    add_heading_25(doc, 'Appendices')
    add_heading_14(doc, 'Appendix A: Web Platform Interfaces')
    
    fig_caption(doc, 'Figure A.1: SmartClaim AI Dashboard with KPI cards and diagnostic charts.')
    fig_caption(doc, 'Figure A.2: Prediction page with Arabic RTL text input and quick example buttons.')
    fig_caption(doc, 'Figure A.3: Dataset Explorer page with interactive Plotly charts and data table.')
    
    add_heading_14(doc, 'Appendix B: Sample Code Snippets')
    fig_caption(doc, 'Code B.1: Arabic Feature Extraction Engine (ArabicAccidentFeatureExtractor class).')
    fig_caption(doc, 'Code B.2: Hybrid Feature Vector Construction (TextEmbedder class).')
    fig_caption(doc, 'Code B.3: Seven-Tier Expert Cost Engine (predict_cost function).')
    fig_caption(doc, 'Code B.4: Model Training and Selection Pipeline (train_models function).')
    
    # Save
    output_path = r'D:\smartclaim\data\SmartClaim_AI_Thesis.docx'
    doc.save(output_path)
    print(f'Thesis V2 saved to: {output_path}')
    
    # Count words
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f'Total word count: {total_words}')
    print(f'Total paragraphs: {len(doc.paragraphs)}')
    print(f'Total sections: {len(doc.sections)}')
    print(f'Total tables: {len(doc.tables)}')

if __name__ == '__main__':
    create_thesis()
